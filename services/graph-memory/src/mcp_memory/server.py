# -*- coding: utf-8 -*-
"""
MCP Memory Server - Serveur principal.

Expose tous les outils MCP via Streamable HTTP avec FastMCP.
"""

import os
import sys
import json
import uuid
import base64
import argparse
from typing import Annotated, Optional, List, Dict, Any

import uvicorn
from dotenv import load_dotenv
from pydantic import Field

# Charger .env avant les imports qui en dépendent
load_dotenv()

from mcp.server.fastmcp import FastMCP, Context

from .config import get_settings
from .auth.middleware import AuthMiddleware, LoggingMiddleware, StaticFilesMiddleware
from .auth.context import check_memory_access, check_write_permission, check_admin_permission, get_allowed_memory_ids, current_auth, DENY_ALL
from .core.validators import validate_memory_id, validate_filename, validate_document_size, validate_entity_name, validate_backup_id as validate_backup_id_format, check_bootstrap_key_safety
from .core.egress import redact_proxy_secrets


# =============================================================================
# Initialisation
# =============================================================================

settings = get_settings()

# Créer l'instance FastMCP
# host="0.0.0.0" pour accepter les connexions externes (reverse proxy, Docker)
mcp = FastMCP(
    name=settings.mcp_server_name,
    host=settings.mcp_server_host,
    port=settings.mcp_server_port,
)


# =============================================================================
# Helpers - Services (lazy-loaded)
# =============================================================================

_graph_service = None
_storage_service = None
_extractor_service = None
_token_manager = None
_embedding_service = None
_chunker = None
_vector_store = None


def get_graph():
    """Lazy-load GraphService."""
    global _graph_service
    if _graph_service is None:
        from .core.graph import get_graph_service
        _graph_service = get_graph_service()
    return _graph_service


def get_storage():
    """Lazy-load StorageService."""
    global _storage_service
    if _storage_service is None:
        from .core.storage import get_storage_service
        _storage_service = get_storage_service()
    return _storage_service


def get_extractor():
    """Lazy-load ExtractorService."""
    global _extractor_service
    if _extractor_service is None:
        from .core.extractor import get_extractor_service
        _extractor_service = get_extractor_service()
    return _extractor_service


def get_tokens():
    """Lazy-load TokenManager."""
    global _token_manager
    if _token_manager is None:
        from .auth.token_manager import get_token_manager
        _token_manager = get_token_manager()
    return _token_manager


def get_embedder():
    """Lazy-load EmbeddingService."""
    global _embedding_service
    if _embedding_service is None:
        from .core.embedder import get_embedding_service
        _embedding_service = get_embedding_service()
    return _embedding_service


def get_chunker():
    """Lazy-load SemanticChunker."""
    global _chunker
    if _chunker is None:
        from .core.chunker import get_chunker as _get_chunker
        _chunker = _get_chunker()
    return _chunker


def get_vector_store():
    """Lazy-load VectorStoreService."""
    global _vector_store
    if _vector_store is None:
        from .core.vector_store import get_vector_store as _get_vs
        _vector_store = _get_vs()
    return _vector_store


_backup_service = None

def get_backup():
    """Lazy-load BackupService."""
    global _backup_service
    if _backup_service is None:
        from .core.backup import get_backup_service
        _backup_service = get_backup_service()
    return _backup_service


# =============================================================================
# OUTILS MCP - Gestion des Mémoires
# =============================================================================

@mcp.tool(description="Create an isolated graph-memory namespace.")
async def memory_create(
    memory_id: Annotated[str, Field(description="Unique memory identifier (for example, 'quoteflow-legal')")],
    name: Annotated[str, Field(description="Human-readable memory name")],
    ontology: Annotated[str, Field(description="Ontology to use (for example, legal, cloud, managed-services, technical, presales)")],
    description: Annotated[Optional[str], Field(default=None, description="Optional memory description")] = None
) -> dict:
    """
    Crée une nouvelle mémoire (namespace isolé).
    
    L'ontologie est OBLIGATOIRE et copiée sur S3 pour persistance et versioning.
    
    Args:
        memory_id: Identifiant unique (ex: "quoteflow-legal")
        name: Nom lisible de la mémoire
        ontology: Nom de l'ontologie à utiliser (OBLIGATOIRE: legal, cloud, managed-services, technical)
        description: Description optionnelle
        
    Returns:
        Informations sur la mémoire créée
    """
    try:
        # Sécurité v2.1.0 : valider memory_id (anti injection Cypher/S3)
        validate_memory_id(memory_id)
        
        # Vérifier la permission d'écriture
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        # Pour memory_create, on vérifie l'accès SAUF si le token a des memory_ids
        # restreints et que la nouvelle mémoire n'y est pas encore (elle sera auto-ajoutée)
        auth = current_auth.get()
        _auto_add_to_token = False
        if auth and auth.get("type") == "token":
            allowed = auth.get("memory_ids", [])
            if allowed and memory_id not in allowed:
                # Token restreint, la mémoire n'est pas dans la liste
                # → On autorise la création et on l'ajoutera automatiquement au token
                _auto_add_to_token = True
                print(f"🔑 [Auth] memory_create: '{memory_id}' will be added to the token for '{auth.get('client_name')}'", file=sys.stderr)
            elif not allowed:
                # memory_ids vide = accès à tout → pas besoin d'auto-add
                pass
            else:
                # La mémoire est déjà dans la liste → OK
                pass
        else:
            # Pas de token restreint → vérification standard
            access_err = check_memory_access(memory_id)
            if access_err:
                return access_err
        
        # Vérifier que l'ontologie existe et la récupérer
        from .core.ontology import get_ontology_manager
        ontology_manager = get_ontology_manager()
        ontology_data = ontology_manager.get_ontology(ontology)
        
        if not ontology_data:
            available = [o["name"] for o in ontology_manager.list_ontologies()]
            return {
                "status": "error",
                "message": f"Ontology '{ontology}' not found. Available: {available}"
            }
        
        # Stocker l'ontologie sur S3 pour la mémoire
        import yaml
        ontology_yaml = yaml.dump(ontology_data, allow_unicode=True, default_flow_style=False)
        ontology_bytes = ontology_yaml.encode('utf-8')
        
        ontology_s3_result = await get_storage().upload_document(
            memory_id=memory_id,
            filename=f"_ontology_{ontology}.yaml",
            content=ontology_bytes,
            metadata={"type": "ontology", "ontology_name": ontology}
        )
        
        print(f"📝 [Memory] Ontology '{ontology}' stored: {ontology_s3_result['uri']}", file=sys.stderr)
        
        # Créer la mémoire dans le graphe avec l'URI S3 de l'ontologie
        memory = await get_graph().create_memory(
            memory_id=memory_id,
            name=name,
            description=description,
            ontology=ontology,
            ontology_uri=ontology_s3_result["uri"]
        )
        
        # Auto-ajouter la mémoire au token si nécessaire (isolation multi-tenant)
        if _auto_add_to_token and auth and auth.get("token_hash"):
            try:
                await get_tokens().update_token_memories(
                    token_hash=auth["token_hash"],
                    add_memories=[memory_id]
                )
                # Mettre à jour le contexte auth en mémoire pour la session courante
                auth["memory_ids"].append(memory_id)
                current_auth.set(auth)
                print(f"🔑 [Auth] memory_create: '{memory_id}' added to the token for '{auth.get('client_name')}'", file=sys.stderr)
            except Exception as e:
                print(f"⚠️ [Auth] Unable to add '{memory_id}' to the token automatically: {e}", file=sys.stderr)
        
        return {
            "status": "created",
            "memory_id": memory.id,
            "name": memory.name,
            "description": memory.description,
            "ontology": memory.ontology,
            "ontology_uri": ontology_s3_result["uri"]
        }
    except ValueError as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"Creation failed: {str(e)}"}


@mcp.tool(description="Update a memory's name or description.")
async def memory_update(
    memory_id: Annotated[str, Field(description="Identifier of the memory to update")],
    name: Annotated[Optional[str], Field(default=None, description="Nouveau nom (vide = pas de changement)")] = None,
    description: Annotated[Optional[str], Field(default=None, description="Nouvelle description (vide = pas de changement)")] = None,
) -> dict:
    """
    Met à jour les métadonnées d'une mémoire (nom, description).
    
    Seuls les champs fournis sont modifiés. L'ontologie n'est pas modifiable.
    
    Args:
        memory_id: ID de la mémoire à modifier
        name: Nouveau nom (None = pas de changement)
        description: Nouvelle description (None = pas de changement)
        
    Returns:
        Mémoire mise à jour avec ses nouvelles valeurs
    """
    try:
        # Vérifier l'accès à la mémoire + permission write
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        if name is None and description is None:
            return {"status": "error", "message": "Nothing to update. Provide --name and/or --description."}
        
        memory = await get_graph().update_memory(memory_id, name=name, description=description)
        
        if not memory:
            return {"status": "error", "message": f"Memory '{memory_id}' not found"}
        
        return {
            "status": "ok",
            "memory_id": memory.id,
            "name": memory.name,
            "description": memory.description,
            "ontology": memory.ontology,
        }
    except Exception as e:
        return {"status": "error", "message": f"Update failed: {str(e)}"}


@mcp.tool(description="Permanently delete a memory and all of its content.")
async def memory_delete(
    memory_id: Annotated[str, Field(description="Identifier of the memory to delete; this cannot be undone")]
) -> dict:
    """
    Supprime une mémoire et tout son contenu (graphe + S3).
    
    ⚠️ ATTENTION: Cette opération est irréversible !
    Supprime le namespace Neo4j ET tous les fichiers S3 associés.
    
    Args:
        memory_id: ID de la mémoire à supprimer
        
    Returns:
        Statut de la suppression avec détails S3
    """
    try:
        # Vérifier l'accès à la mémoire + permission write
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        # 1. Supprimer la collection Qdrant (couplage strict)
        qdrant_deleted = False
        try:
            qdrant_deleted = await get_vector_store().delete_collection(memory_id)
        except Exception as e:
            print(f"❌ [Qdrant] Error deleting collection for {memory_id}: {e}", file=sys.stderr)
            raise RuntimeError(f"Unable to delete the Qdrant collection (strict coupling): {e}")
        
        # 2. Supprimer tous les fichiers S3 de la mémoire
        s3_result = {"deleted_count": 0, "error_count": 0}
        try:
            s3_result = await get_storage().delete_prefix(f"{memory_id}/")
            print(f"🗑️ [S3] Memory {memory_id} cleanup: deleted {s3_result['deleted_count']} files", file=sys.stderr)
        except Exception as e:
            print(f"⚠️ [S3] S3 cleanup error for {memory_id}: {e}", file=sys.stderr)
        
        # 3. Supprimer du graphe Neo4j
        deleted = await get_graph().delete_memory(memory_id)
        
        if deleted:
            return {
                "status": "deleted",
                "memory_id": memory_id,
                "qdrant_collection_deleted": qdrant_deleted,
                "s3_files_deleted": s3_result.get("deleted_count", 0),
                "s3_errors": s3_result.get("error_count", 0)
            }
        return {"status": "not_found", "memory_id": memory_id}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="List the memories available to the caller.")
async def memory_list() -> dict:
    """
    Liste les mémoires accessibles au client.
    
    Un client non-admin ne voit que les mémoires autorisées par son token.
    Un admin ou un accès localhost voit toutes les mémoires.
    
    Returns:
        Liste des mémoires avec leurs métadonnées
    """
    try:
        memories = await get_graph().list_memories()
        
        # Filtrer par les mémoires autorisées (isolation multi-tenant)
        allowed = get_allowed_memory_ids()
        # P7-4 (ADR-0019): fail-closed — no auth context => deny the listing.
        if allowed is DENY_ALL:
            return {"status": "error", "message": "Authentication required"}
        if allowed is not None and len(allowed) > 0:
            # Token restreint : ne montrer que les mémoires autorisées
            memories = [m for m in memories if m.id in allowed]
        # allowed is None → admin/bootstrap/localhost → pas de filtrage
        # allowed == [] → token sans restriction → pas de filtrage
        
        return {
            "status": "ok",
            "count": len(memories),
            "memories": [
                {
                    "id": m.id,
                    "name": m.name,
                    "description": m.description,
                    "ontology": m.ontology,
                    "created_at": m.created_at.isoformat() if m.created_at else None
                }
                for m in memories
            ]
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Return statistics for a memory.")
async def memory_stats(
    memory_id: Annotated[str, Field(description="Memory identifier")]
) -> dict:
    """
    Récupère les statistiques d'une mémoire.
    
    Args:
        memory_id: ID de la mémoire
        
    Returns:
        Statistiques (documents, entités, relations, top entités)
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        stats = await get_graph().get_memory_stats(memory_id)
        return {
            "status": "ok",
            "memory_id": memory_id,
            "document_count": stats.document_count,
            "entity_count": stats.entity_count,
            "relation_count": stats.relation_count,
            "top_entities": stats.top_entities
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Ingestion de Documents
# =============================================================================

@mcp.tool(description="Ingest a document into a memory synchronously.")
async def memory_ingest(
    memory_id: Annotated[str, Field(description="Target memory identifier")],
    content_base64: Annotated[str, Field(description="Base64-encoded document content")],
    filename: Annotated[str, Field(description="Filename (for example, 'contract.pdf' or 'notes.md')")],
    metadata: Annotated[Optional[Dict[str, Any]], Field(default=None, description="Optional free-form key/value metadata")] = None,
    force: Annotated[bool, Field(default=False, description="Reingest even when the document already exists (SHA-256 deduplication)")] = False,
    source_path: Annotated[Optional[str], Field(default=None, description="Original source path (for example, 'legal/contracts/terms.pdf')")] = None,
    source_modified_at: Annotated[Optional[str], Field(default=None, description="Date de modification source ISO 8601 (ex: '2026-01-15T10:30:00')")] = None,
    ctx: Optional[Context] = None
) -> dict:
    """
    Ingère un document dans une mémoire.
    
    Le document est:
    1. Stocké sur S3
    2. Analysé par le LLM pour extraire entités/relations
    3. Les entités et relations sont ajoutées au graphe
    
    Métadonnées enrichies stockées sur le nœud Document :
    - hash SHA-256 (déduplication)
    - taille en bytes, longueur du texte extrait
    - type de fichier (extension)
    - chemin source et date de modification source (si fournis)
    - stats d'extraction (entités, relations, chunks)
    
    Args:
        memory_id: ID de la mémoire cible
        content_base64: Contenu du document encodé en base64
        filename: Nom du fichier
        metadata: Métadonnées additionnelles (optionnel)
        force: Si True, réingère même si le document existe déjà
        source_path: Chemin complet d'origine du fichier (ex: "legal/contracts/CGA.pdf")
        source_modified_at: Date de dernière modification du fichier source (ISO 8601, ex: "2026-01-15T10:30:00")
        
    Returns:
        Résultat de l'ingestion avec statistiques
    """
    try:
        # Sécurité v2.1.0 : valider les entrées (anti injection S3/path traversal)
        validate_memory_id(memory_id)
        filename = validate_filename(filename)

        # Vérifier l'accès à la mémoire + permission write
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err

        # Décoder le contenu (libérer content_base64 ensuite — peut être volumineux)
        content = base64.b64decode(content_base64)
        content_size = len(content)
        del content_base64
        if ctx:
            try:
                await ctx.info(f"📦 Decoding: {content_size} bytes ({filename})")
            except Exception:
                pass

        # Sécurité v2.1.0 : limite de taille document (anti DoS)
        validate_document_size(content, settings.max_document_size_bytes)

        # Calculer le hash (déduplication + clé de changement)
        doc_hash = get_storage().compute_hash(content)

        # Déduplication historique par hash (sémantique synchrone conservée)
        existing = await get_graph().get_document_by_hash(memory_id, doc_hash)
        if existing and not force:
            return {
                "status": "already_exists",
                "document_id": existing.id,
                "filename": existing.filename,
                "message": "Document already ingested (use force=true to reingest)"
            }
        replace_doc_id = existing.id if (existing and force) else None

        # Relai vers le pipeline factorisé (commun avec l'ingestion asynchrone).
        # Le pipeline finalise ingestion_status="succeeded" après Qdrant.
        async def _progress_cb(step: str, percent: int, extra: dict):
            if ctx:
                try:
                    await ctx.info(extra.get("message", step))
                except Exception:
                    pass

        from .core.ingest_pipeline import run_ingest_pipeline
        return await run_ingest_pipeline(
            memory_id=memory_id,
            content=content,
            filename=filename,
            doc_hash=doc_hash,
            metadata=metadata,
            source_path=source_path,
            source_modified_at=source_modified_at,
            last_ingest_job_id=None,
            replace_doc_id=replace_doc_id,
            progress_cb=_progress_cb,
        )

    except Exception as e:
        print(f"❌ [Ingest] Error: {e}", file=sys.stderr)
        return {"status": "error", "message": str(e)}


def _extract_text(content: bytes, filename: str) -> Optional[str]:
    """
    Extrait le texte d'un document.
    
    Formats supportés: txt, md, html, docx, pdf, csv
    """
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    try:
        # Texte brut et Markdown
        if ext in ('txt', 'md'):
            return content.decode('utf-8', errors='ignore')
        
        # HTML
        elif ext in ('html', 'htm'):
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'html.parser')
            # Supprimer scripts et styles
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
            return text
        
        # PDF
        elif ext == 'pdf':
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        
        # DOCX (Word)
        elif ext == 'docx':
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content))
            
            text_parts = []
            
            # Extraire les paragraphes
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        
        # CSV
        elif ext == 'csv':
            import csv
            from io import StringIO
            
            # Décoder le contenu
            text_content = content.decode('utf-8', errors='ignore')
            reader = csv.reader(StringIO(text_content))
            
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
            
            return "\n".join(rows)
        
        else:
            # Tenter de décoder comme texte (fallback)
            return content.decode('utf-8', errors='ignore')
            
    except Exception as e:
        print(f"⚠️ [Extract] Text extraction error ({ext}): {e}", file=sys.stderr)
        return None


# =============================================================================
# OUTILS MCP - Ingestion asynchrone
# =============================================================================

@mcp.tool(description="Queue a document for asynchronous ingestion.")
async def memory_ingest_async(
    memory_id: Annotated[str, Field(description="Target memory identifier")],
    content_base64: Annotated[str, Field(description="Base64-encoded document content")],
    filename: Annotated[str, Field(description="Filename (for example, 'contract.pdf' or 'notes.md')")],
    source_path: Annotated[str, Field(description="Stable source-path business key (for example, 'legal/contracts/terms.pdf'); required")],
    sha256: Annotated[str, Field(description="SHA-256 checksum of the decoded content in hexadecimal; required for integrity and change detection")],
    metadata: Annotated[Optional[Dict[str, Any]], Field(default=None, description="Optional metadata")] = None,
    source_modified_at: Annotated[Optional[str], Field(default=None, description="Date de modification source ISO 8601")] = None,
    job_id: Annotated[Optional[str], Field(default=None, description="Optional client-supplied job ID for idempotency; generated when omitted")] = None,
    replace_existing: Annotated[bool, Field(default=False, description="Replace an existing document when its checksum changed; defaults to false and returns changed_skipped")] = False,
) -> dict:
    """
    Soumet un document à l'ingestion ASYNCHRONE et rend la main immédiatement.

    L'extraction LLM + embeddings se déroulent en tâche de fond (un worker par
    mémoire). Suivi via `ingest_job_status`, listing via `ingest_job_list`,
    annulation via `ingest_job_cancel`.

    Idempotence par `source_path` (clé métier) + `sha256` (détecteur de changement) :
    - source_path + checksum déjà ingérés avec succès → `skipped` (immédiat, sans job)
    - checksum différent + replace_existing=false → `changed_skipped` (immédiat)
    - checksum différent + replace_existing=true → remplacement propre
    - nouveau → `queued`/`running`

    Returns:
        {job_id?, status, document_id?, message} — réponse immédiate.
    """
    try:
        validate_memory_id(memory_id)
        filename = validate_filename(filename)

        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err

        if not source_path or not source_path.strip():
            return {"status": "error", "message": "source_path is required as a stable business key."}
        if not sha256 or not sha256.strip():
            return {"status": "error", "message": "sha256 est obligatoire (checksum du contenu)."}

        content = base64.b64decode(content_base64)
        del content_base64
        validate_document_size(content, settings.max_document_size_bytes)

        # Garde d'intégrité : le checksum fourni doit correspondre au contenu reçu
        computed = get_storage().compute_hash(content)
        provided = sha256.strip().lower()
        if computed.lower() != provided:
            return {
                "status": "error",
                "message": f"Invalid checksum: supplied SHA-256 ({provided[:12]}…) does not match received content ({computed[:12]}…).",
            }

        requested_by = ""
        try:
            auth = current_auth.get()
            requested_by = getattr(auth, "client_name", "") or getattr(auth, "name", "") or ""
        except Exception:
            pass

        from .core.ingest_queue import get_ingest_queue
        return await get_ingest_queue().submit(
            memory_id=memory_id,
            content=content,
            filename=filename,
            sha256=computed,
            source_path=source_path,
            replace_existing=replace_existing,
            metadata=metadata,
            source_modified_at=source_modified_at,
            requested_by=requested_by,
            job_id=job_id,
        )
    except Exception as e:
        print(f"❌ [IngestAsync] Error: {e}", file=sys.stderr)
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Queue a batch of documents for asynchronous ingestion.")
async def memory_ingest_batch_async(
    memory_id: Annotated[str, Field(description="Target memory identifier")],
    documents: Annotated[List[Dict[str, Any]], Field(description="Documents, each containing {content_base64, filename, source_path, sha256, metadata?, source_modified_at?}")],
    replace_existing: Annotated[bool, Field(default=False, description="Replace documents with changed checksums across the entire batch")] = False,
) -> dict:
    """
    Soumet un LOT de documents à l'ingestion asynchrone.

    Chaque document est résolu indépendamment (source_path + sha256) puis mis en
    file. Un `batch_id` commun permet de suivre l'avancement agrégé via ce même
    outil (réappel) ou `ingest_job_list(memory_id, ...)`.

    Returns:
        {batch_id, total, counts{queued/running/succeeded/failed/skipped/...},
         errors[], items[]}
    """
    try:
        validate_memory_id(memory_id)
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err

        if not documents:
            return {"status": "error", "message": "No document provided."}

        requested_by = ""
        try:
            auth = current_auth.get()
            requested_by = getattr(auth, "client_name", "") or getattr(auth, "name", "") or ""
        except Exception:
            pass

        import uuid as _uuid
        batch_id = f"batch_{_uuid.uuid4().hex}"

        from .core.ingest_queue import get_ingest_queue
        queue = get_ingest_queue()

        items = []
        counts = {k: 0 for k in ("queued", "running", "succeeded", "failed", "skipped", "changed_skipped", "error", "queue_full")}
        errors = []

        for idx, doc in enumerate(documents):
            try:
                fname = validate_filename(doc.get("filename", ""))
                sp = doc.get("source_path", "")
                sha = (doc.get("sha256") or "").strip()
                cb64 = doc.get("content_base64", "")
                if not sp or not sha or not cb64:
                    raise ValueError("content_base64, filename, source_path, and sha256 are required for each document")

                content = base64.b64decode(cb64)
                validate_document_size(content, settings.max_document_size_bytes)
                computed = get_storage().compute_hash(content)
                if computed.lower() != sha.lower():
                    raise ValueError(f"Invalid checksum for {fname}")

                res = await queue.submit(
                    memory_id=memory_id,
                    content=content,
                    filename=fname,
                    sha256=computed,
                    source_path=sp,
                    replace_existing=replace_existing,
                    metadata=doc.get("metadata"),
                    source_modified_at=doc.get("source_modified_at"),
                    requested_by=requested_by,
                    batch_id=batch_id,
                )
            except Exception as item_err:
                res = {"status": "error", "message": str(item_err), "source_path": doc.get("source_path")}
                errors.append({"source_path": doc.get("source_path"), "filename": doc.get("filename"), "error": str(item_err)})

            st = res.get("status", "error")
            counts[st] = counts.get(st, 0) + 1
            items.append({"index": idx, "source_path": doc.get("source_path"), "job_id": res.get("job_id"), "status": st})

        return {
            "status": "ok",
            "batch_id": batch_id,
            "memory_id": memory_id,
            "total": len(documents),
            "counts": counts,
            "errors": errors,
            "items": items,
            "message": "Batch submitted. Track it by batch_id with ingest_job_list or call again.",
        }
    except Exception as e:
        print(f"❌ [IngestBatchAsync] Error: {e}", file=sys.stderr)
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Return the status of an asynchronous ingestion job.")
async def ingest_job_status(
    job_id: Annotated[str, Field(description="Job ID returned by memory_ingest_async")]
) -> dict:
    """
    Consulte l'état d'un job d'ingestion asynchrone.

    Returns:
        status, current_step, progress_percent, created_entities,
        created_relations, started_at/updated_at/finished_at, error éventuelle.
    """
    try:
        from .core.ingest_queue import get_ingest_queue
        result = await get_ingest_queue().get_job(job_id)
        # Contrôle d'accès si le job est connu
        mem = result.get("memory_id")
        if mem:
            access_err = check_memory_access(mem)
            if access_err:
                return access_err
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="List asynchronous ingestion jobs for a memory.")
async def ingest_job_list(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    status: Annotated[Optional[str], Field(default=None, description="Filtre par statut (queued|running|succeeded|failed|cancelled|skipped|changed_skipped)")] = None,
    source_path: Annotated[Optional[str], Field(default=None, description="Filter by source_path when resuming after a client timeout")] = None,
    batch_id: Annotated[Optional[str], Field(default=None, description="Filtre par batch_id")] = None,
) -> dict:
    """
    Liste les jobs d'ingestion d'une mémoire (reprise après timeout client).

    Permet de retrouver un job par `source_path` après une coupure réseau.
    Note : l'historique des jobs est in-memory best-effort (perdu au redémarrage
    du conteneur) ; l'état d'ingestion durable reste lisible via `document_list`.
    """
    try:
        validate_memory_id(memory_id)
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        from .core.ingest_queue import get_ingest_queue
        return await get_ingest_queue().list_jobs(memory_id, status=status, source_path=source_path, batch_id=batch_id)
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Request cancellation of an asynchronous ingestion job.")
async def ingest_job_cancel(
    job_id: Annotated[str, Field(description="Identifier of the job to cancel")]
) -> dict:
    """
    Annule un job d'ingestion (best-effort, sans corrompre le graphe).

    - Job en attente : retiré immédiatement de la file (`cancelled`).
    - Job en cours : annulation coopérative à la prochaine frontière de phase ;
      les écritures partielles sont nettoyées (aucun orphelin).
    """
    try:
        from .core.ingest_queue import get_ingest_queue
        queue = get_ingest_queue()
        # Contrôle d'accès via la mémoire du job
        info = await queue.get_job(job_id)
        mem = info.get("memory_id")
        if mem:
            access_err = check_memory_access(mem)
            if access_err:
                return access_err
            write_err = check_write_permission()
            if write_err:
                return write_err
        return await queue.cancel(job_id)
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Recherche
# =============================================================================

@mcp.tool(description="Search a memory using graph-first retrieval.")
async def memory_search(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    query: Annotated[str, Field(description="Free-text search query")],
    limit: Annotated[int, Field(default=10, description="Maximum number of results; defaults to 10")] = 10
) -> dict:
    """
    Recherche dans une mémoire (graph-first).
    
    Recherche les entités et documents correspondant à la requête.
    Utilise principalement le graphe, pas de RAG vectoriel.
    
    Args:
        memory_id: ID de la mémoire
        query: Requête de recherche
        limit: Nombre max de résultats
        
    Returns:
        Entités trouvées avec leurs documents liés
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # Recherche d'entités
        entities = await get_graph().search_entities(memory_id, search_query=query, limit=limit)
        
        # Pour chaque entité, récupérer le contexte complet
        results = []
        for entity in entities:
            context = await get_graph().get_entity_context(
                memory_id, entity["name"], depth=1
            )
            results.append({
                "entity": entity,
                "documents": context.documents,
                "related_entities": context.related_entities
            })
        
        return {
            "status": "ok",
            "query": query,
            "memory_id": memory_id,
            "result_count": len(results),
            "results": results
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Answer a question using the graph and document context.")
async def question_answer(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    question: Annotated[str, Field(description="Natural-language question")],
    limit: Annotated[int, Field(default=10, description="Maximum number of entities to search; defaults to 10")] = 10
) -> dict:
    """
    Pose une question sur une mémoire et obtient une réponse basée sur le graphe.
    
    Utilise le graphe de connaissances pour répondre à la question.
    Recherche les entités pertinentes puis génère une réponse avec le LLM.
    
    Args:
        memory_id: ID de la mémoire
        question: Question en langage naturel
        limit: Nombre max d'entités à rechercher (défaut: 10)
        
    Returns:
        Réponse générée avec les entités liées
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # 1. Rechercher les entités pertinentes dans le graphe
        print(f"🔎 [Q&A] Recherche graphe: memory={memory_id}, question='{question}', limit={limit}", file=sys.stderr)
        entities = await get_graph().search_entities(memory_id, search_query=question, limit=limit)
        
        if entities:
            entity_summary = ", ".join(f"{e['name']} ({e.get('type','?')})" for e in entities)
            print(f"📊 [Q&A] Graph: found {len(entities)} entities → {entity_summary}", file=sys.stderr)
        else:
            print("📊 [Q&A] Graph: found 0 entities → RAG-only fallback", file=sys.stderr)
        
        # 2. Récupérer le contexte de chaque entité + documents sources
        context_parts = []
        entity_names = []
        source_documents = {}  # doc_id -> {filename, id}

        for entity in entities:
            entity_names.append(entity["name"])
            ctx = await get_graph().get_entity_context(memory_id, entity["name"], depth=1)

            # Collecter les documents sources et les associer à l'entité
            entity_doc_names = []
            for doc in ctx.documents:
                if isinstance(doc, dict):
                    doc_id = doc.get('id', '')
                    doc_filename = doc.get('filename', doc_id)
                    if doc_id:
                        if doc_id not in source_documents:
                            source_documents[doc_id] = {
                                "id": doc_id,
                                "filename": doc_filename,
                            }
                        entity_doc_names.append(doc_filename)

            # Construire le contexte texte AVEC le document source
            doc_ref = f" [Source: {', '.join(entity_doc_names)}]" if entity_doc_names else ""
            ctx_text = f"- {entity['name']} ({entity.get('type', '?')}){doc_ref}"
            if entity.get('description'):
                ctx_text += f": {entity['description']}"

            for rel in ctx.relations:
                ctx_text += f"\n  → {rel.get('type', 'RELATED_TO')}: {rel.get('description', '')}"

            related = [r['name'] for r in ctx.related_entities]
            if related:
                ctx_text += f"\n  Related to: {', '.join(related)}"

            context_parts.append(ctx_text)

        # 3. === RAG vectoriel : Graph-Guided si entités trouvées, sinon RAG-only ===
        rag_context_parts = []
        rag_chunks_used = 0
        rag_mode = "graph-guided" if entities else "rag-only"
        try:
            # Collecter les doc_ids identifiés par le graphe (vide si aucune entité)
            graph_doc_ids = list(source_documents.keys())

            # Vectoriser la question
            query_embedding = await get_embedder().embed_query(question)

            # Recherche Qdrant :
            # - Graph-Guided : filtrée par les documents identifiés par le graphe
            # - RAG-only : recherche sur TOUS les chunks de la mémoire (fallback)
            score_threshold = settings.rag_score_threshold
            chunk_limit = settings.rag_chunk_limit
            
            chunk_results = await get_vector_store().search(
                memory_id=memory_id,
                query_embedding=query_embedding,
                doc_ids=graph_doc_ids if graph_doc_ids else None,
                limit=chunk_limit
            )

            # Sauver tous les résultats avant filtrage (pour diagnostic)
            all_chunk_results = list(chunk_results)
            
            # Filtrer par seuil de score (en dessous = non pertinent)
            total_before = len(chunk_results)
            chunk_results = [cr for cr in chunk_results if cr.score >= score_threshold]
            filtered_out = total_before - len(chunk_results)

            # Construire le contexte RAG (chunks pertinents)
            for cr in chunk_results:
                rag_context_parts.append(cr.context_text)
                rag_chunks_used += 1
                # Ajouter les docs trouvés par RAG au source_documents
                if cr.chunk.doc_id and cr.chunk.doc_id not in source_documents:
                    source_documents[cr.chunk.doc_id] = {
                        "id": cr.chunk.doc_id,
                        "filename": cr.chunk.filename or "?",
                    }

            print(f"🔍 [Q&A] RAG ({rag_mode}): {rag_chunks_used} chunks retenus"
                  f" (threshold={score_threshold}, filtered {filtered_out} of {total_before})"
                  f"{f' | graph-guided: {len(graph_doc_ids)} docs' if graph_doc_ids else ' | tous documents'}", 
                  file=sys.stderr)
            
            # Log détaillé : score + section + aperçu texte de chaque chunk RETENU
            for i, cr in enumerate(chunk_results):
                section = cr.chunk.section_title or cr.chunk.article_number or "—"
                preview = cr.chunk.text[:80].replace('\n', ' ').strip()
                print(f"   📎 [{i+1}] score={cr.score:.4f} ✅ | {section} | \"{preview}...\"", file=sys.stderr)
            
            # Log des chunks FILTRÉS (sous le seuil) — diagnostic de pertinence RAG
            if filtered_out > 0:
                # Recalculer les chunks filtrés pour le log
                filtered_chunks = [cr for cr in all_chunk_results if cr.score < score_threshold]
                for i, cr in enumerate(filtered_chunks[:5]):  # Max 5 pour ne pas surcharger
                    section = cr.chunk.section_title or cr.chunk.article_number or "—"
                    preview = cr.chunk.text[:60].replace('\n', ' ').strip()
                    print(f"   📎 [F{i+1}] score={cr.score:.4f} ❌ | {section} | \"{preview}...\"", file=sys.stderr)

        except Exception as e:
            print(f"⚠️ [Q&A] Vector RAG error: {e}", file=sys.stderr)
            # On continue avec le contexte graphe seul

        # Si ni le graphe ni le RAG n'ont trouvé quoi que ce soit → pas de contexte
        if not entities and rag_chunks_used == 0:
            return {
                "status": "ok",
                "answer": "I could not find relevant information in this memory to answer your question.",
                "entities": [],
                "rag_chunks_used": 0,
                "source_documents": []
            }
        
        # 4. Construire la liste des documents pour le prompt
        doc_list = "\n".join(
            f"  - {doc['filename']}" for doc in source_documents.values()
        )
        
        # 5. Assembler le contexte final (graphe + RAG)
        graph_context = "\n".join(context_parts)
        rag_context = "\n\n".join(rag_context_parts) if rag_context_parts else ""
        
        # 6. Générer la réponse avec le LLM
        graph_ctx_len = len(graph_context) if graph_context else 0
        rag_ctx_len = len(rag_context) if rag_context else 0
        doc_count = len(source_documents)
        print(f"📝 [Q&A] Contexte LLM: graphe={graph_ctx_len} chars, RAG={rag_ctx_len} chars, docs={doc_count}", file=sys.stderr)
        
        prompt = f"""You answer questions using a knowledge graph and document excerpts.

Available source documents:
{doc_list}

=== CONTEXT 1: Knowledge graph (entities and relations) ===
{graph_context}

=== CONTEXT 2: Relevant document excerpts (vector RAG) ===
{rag_context if rag_context else "(no additional excerpts)"}

User question: {question}

INSTRUCTIONS:
- Answer concisely and precisely using only the supplied context.
- Prefer document excerpts (CONTEXT 2) for factual details and citations.
- Use the graph (CONTEXT 1) for overview and relationships between concepts.
- Cite the source document for every factual claim.
- Name every source when information comes from multiple documents.
- Clearly say when the context is insufficient for a complete answer.
- Structure the response with Markdown.
"""
        
        answer = await get_extractor().generate_answer(prompt)
        
        return {
            "status": "ok",
            "answer": answer,
            "entities": entity_names,
            "rag_chunks_used": rag_chunks_used,
            "source_documents": list(source_documents.values()),
            "context_used": graph_context
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Query structured memory data without LLM generation.")
async def memory_query(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    query: Annotated[str, Field(description="Natural-language query")],
    limit: Annotated[int, Field(default=10, description="Maximum number of entities to search; defaults to 10")] = 10
) -> dict:
    """
    Interroge une mémoire et retourne les données structurées SANS génération LLM.
    
    Effectue la même recherche que question_answer (graphe + RAG vectoriel)
    mais retourne les données brutes structurées au lieu de crafter une réponse.
    Idéal pour les agents IA qui veulent construire leur propre réponse.
    
    Pipeline :
    1. Recherche d'entités dans le graphe (fulltext + CONTAINS)
    2. Récupération du contexte de chaque entité (voisins, relations, documents)
    3. Recherche RAG vectorielle (graph-guided ou rag-only)
    4. Retour des données structurées (pas d'appel LLM)
    
    Args:
        memory_id: ID de la mémoire
        query: Requête en langage naturel
        limit: Nombre max d'entités à rechercher (défaut: 10)
        
    Returns:
        Données structurées : entités, relations, chunks RAG, documents sources, stats
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # 1. Rechercher les entités pertinentes dans le graphe
        print(f"🔎 [Query] Recherche graphe: memory={memory_id}, query='{query}', limit={limit}", file=sys.stderr)
        entities = await get_graph().search_entities(memory_id, search_query=query, limit=limit)
        
        if entities:
            entity_summary = ", ".join(f"{e['name']} ({e.get('type','?')})" for e in entities)
            print(f"📊 [Query] Graph: found {len(entities)} entities → {entity_summary}", file=sys.stderr)
        else:
            print("📊 [Query] Graph: found 0 entities → RAG-only fallback", file=sys.stderr)
        
        # 2. Récupérer le contexte de chaque entité + documents sources
        enriched_entities = []
        source_documents = {}  # doc_id -> {filename, id}
        meta = {}  # doc_id -> métadonnées enrichies (source_path, repo_path…), source de vérité unique

        for entity in entities:
            ctx = await get_graph().get_entity_context(memory_id, entity["name"], depth=1)

            # Collecter les documents sources
            entity_docs = []
            for doc in ctx.documents:
                if isinstance(doc, dict):
                    doc_id = doc.get('id', '')
                    doc_filename = doc.get('filename', doc_id)
                    if doc_id:
                        # get_entity_context renvoie déjà les docs enrichis (source_path…) :
                        # on alimente la map meta SANS requête supplémentaire.
                        meta.setdefault(doc_id, doc)
                        if doc_id not in source_documents:
                            source_documents[doc_id] = {
                                "id": doc_id,
                                "filename": doc_filename,
                            }
                        entity_docs.append(doc_filename)
            
            # Construire l'entité enrichie
            enriched_entity = {
                "name": entity["name"],
                "type": entity.get("type", "?"),
                "description": entity.get("description", ""),
                "source_documents": entity_docs,
                "relations": [
                    {
                        "type": rel.get("type", "RELATED_TO"),
                        "target": rel.get("target", rel.get("to", "?")),
                        "description": rel.get("description", ""),
                    }
                    for rel in ctx.relations
                ],
                "related_entities": [
                    {
                        "name": r.get("name", r) if isinstance(r, dict) else str(r),
                        "type": r.get("type", "?") if isinstance(r, dict) else "?",
                    }
                    for r in ctx.related_entities
                ],
            }
            enriched_entities.append(enriched_entity)
        
        # 3. RAG vectoriel : Graph-Guided si entités, sinon RAG-only
        rag_chunks = []
        rag_mode = "graph-guided" if entities else "rag-only"
        rag_chunks_filtered = 0
        retained = []  # défini hors du try : utilisé plus bas même si le RAG échoue

        try:
            graph_doc_ids = list(source_documents.keys())
            query_embedding = await get_embedder().embed_query(query)
            
            score_threshold = settings.rag_score_threshold
            chunk_limit = settings.rag_chunk_limit
            
            chunk_results = await get_vector_store().search(
                memory_id=memory_id,
                query_embedding=query_embedding,
                doc_ids=graph_doc_ids if graph_doc_ids else None,
                limit=chunk_limit
            )
            
            total_before = len(chunk_results)
            retained = [cr for cr in chunk_results if cr.score >= score_threshold]
            rag_chunks_filtered = total_before - len(retained)
            
            for cr in retained:
                rag_chunks.append({
                    "text": cr.chunk.text,
                    "score": round(cr.score, 4),
                    "doc_id": cr.chunk.doc_id or "",
                    "filename": cr.chunk.filename or "?",
                    "section_title": cr.chunk.section_title or "",
                    "article_number": cr.chunk.article_number or "",
                    "chunk_index": cr.chunk.index if hasattr(cr.chunk, 'index') else 0,
                })
                # Ajouter les docs trouvés par RAG
                if cr.chunk.doc_id and cr.chunk.doc_id not in source_documents:
                    source_documents[cr.chunk.doc_id] = {
                        "id": cr.chunk.doc_id,
                        "filename": cr.chunk.filename or "?",
                    }
            
            print(f"🔍 [Query] RAG ({rag_mode}): {len(retained)} chunks retenus"
                  f" (threshold={score_threshold}, filtered {rag_chunks_filtered} of {total_before})",
                  file=sys.stderr)
        
        except Exception as e:
            print(f"⚠️ [Query] Vector RAG error: {e}", file=sys.stderr)

        # 3bis. Enrichissement source_path/repo_path par jointure graphe (rétroactif, pas de Qdrant)
        # meta est déjà alimentée par les docs du graphe (étape 2) : on ne complète QUE les
        # doc_ids issus du RAG absents de meta (évite un gros IN sur une entité très fréquente).
        missing = [cr.chunk.doc_id for cr in retained
                   if cr.chunk.doc_id and cr.chunk.doc_id not in meta]
        if missing:
            try:
                meta.update(await get_graph().get_documents_meta(memory_id, list(set(missing))))
            except Exception as e:
                print(f"⚠️ [Query] source_path enrichment failed: {e}", file=sys.stderr)
        print(f"🧭 [Query] {len(meta)} documents enrichis (source_path)", file=sys.stderr)

        def _doc_fields(m):
            """Contrat commun de métadonnées document (hash ET sha256 conservés)."""
            return {
                "uri": m.get("uri"),
                "source_path": m.get("source_path"),
                "repo_path": m.get("repo_path"),
                "hash": m.get("hash"),
                "sha256": m.get("sha256"),
                "ingestion_status": m.get("ingestion_status", "unknown"),
                "chunk_count": m.get("chunk_count", 0),
                "last_ingest_job_id": m.get("last_ingest_job_id"),
            }

        # Enrichir source_documents (contrat complet) — la valeur de meta fait autorité
        enriched_sources = [{**base, **_doc_fields(meta.get(doc_id, {}))}
                            for doc_id, base in source_documents.items()]

        # Enrichir rag_chunks (source_path + repo_path suffisent pour ouvrir le fichier Git)
        for ch in rag_chunks:
            m = meta.get(ch.get("doc_id"), {})
            ch["source_path"] = m.get("source_path")
            ch["repo_path"] = m.get("repo_path")

        # 4. Retourner les données structurées (PAS d'appel LLM)
        return {
            "status": "ok",
            "memory_id": memory_id,
            "query": query,
            "retrieval_mode": rag_mode,
            "entities": enriched_entities,
            "rag_chunks": rag_chunks,
            "source_documents": enriched_sources,
            "stats": {
                "entities_found": len(enriched_entities),
                "rag_chunks_retained": len(rag_chunks),
                "rag_chunks_filtered": rag_chunks_filtered,
                "rag_score_threshold": settings.rag_score_threshold,
                "rag_chunk_limit": settings.rag_chunk_limit,
            },
        }
    
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Return the context surrounding a named entity.")
async def memory_get_context(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    entity_name: Annotated[str, Field(description="Exact name of the entity to explore")],
    depth: Annotated[int, Field(default=1, description="Graph traversal depth; 1 means direct neighbors")] = 1
) -> dict:
    """
    Récupère le contexte complet d'une entité.
    
    Retourne tout ce qu'on sait sur une entité:
    - Documents qui la mentionnent
    - Entités reliées
    - Types de relations
    
    Args:
        memory_id: ID de la mémoire
        entity_name: Nom de l'entité
        depth: Profondeur de traversée (1 = voisins directs)
        
    Returns:
        Contexte complet de l'entité
    """
    try:
        # Sécurité v2.1.0 : valider entity_name (M2)
        validate_entity_name(entity_name)
        
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        context = await get_graph().get_entity_context(
            memory_id, entity_name, depth
        )
        
        return {
            "status": "ok",
            "entity_name": context.entity_name,
            "entity_type": context.entity_type,
            "depth": context.depth,
            "documents": context.documents,
            "related_entities": context.related_entities,
            "relations": context.relations
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Admin / Tokens
# =============================================================================

@mcp.tool(description="Create an access token.")
async def admin_create_token(
    client_name: Annotated[str, Field(description="Client name (for example, 'quoteflow' or 'vela')")],
    permissions: Annotated[Optional[List[str]], Field(default=None, description="Permissions: 'read', 'write', and 'admin'; defaults to ['read', 'write']")] = None,
    memory_ids: Annotated[Optional[List[str]], Field(default=None, description="Allowed memory IDs; empty means all memories")] = None,
    expires_in_days: Annotated[Optional[int], Field(default=None, description="Expiration en jours (optionnel, None = pas d'expiration)")] = None,
    email: Annotated[Optional[str], Field(default=None, description="Token owner's email address")] = None
) -> dict:
    """
    Crée un nouveau token d'accès pour un client.
    
    ⚠️ Le token retourné ne sera affiché qu'une seule fois !
    
    Args:
        client_name: Nom du client (ex: "quoteflow")
        permissions: Permissions ["read", "write", "admin"]
        memory_ids: IDs des mémoires autorisées (vide = toutes)
        expires_in_days: Expiration en jours (optionnel)
        email: Adresse email du propriétaire (optionnel)
        
    Returns:
        Token généré (à conserver précieusement)
    """
    try:
        # Sécurité : permission admin requise
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err
        
        # Valider les permissions
        valid_perms = {"read", "write", "admin"}
        requested_perms = permissions or ["read", "write"]
        invalid = set(requested_perms) - valid_perms
        if invalid:
            return {
                "status": "error",
                "message": f"Permissions invalides: {invalid}. Valides: {sorted(valid_perms)}"
            }
        
        token = await get_tokens().create_token(
            client_name=client_name,
            permissions=permissions or ["read", "write"],
            memory_ids=memory_ids or [],
            expires_in_days=expires_in_days,
            email=email
        )
        
        return {
            "status": "ok",
            "client_name": client_name,
            "email": email,
            "token": token,
            "permissions": permissions or ["read", "write"],
            "memory_ids": memory_ids or [],
            "message": "Save this token now; it will not be shown again."
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="List access tokens.")
async def admin_list_tokens() -> dict:
    """
    Liste tous les tokens actifs.
    
    Note: Les tokens eux-mêmes ne sont pas affichés, seulement leurs métadonnées.
    
    Returns:
        Liste des tokens avec leurs infos
    """
    try:
        # Sécurité : permission admin requise
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err
        
        tokens = await get_tokens().list_tokens()
        
        return {
            "status": "ok",
            "count": len(tokens),
            "tokens": [
                {
                    "client_name": t.client_name,
                    "email": t.email,
                    "permissions": t.permissions,
                    "memory_ids": t.memory_ids,
                    "created_at": t.created_at.isoformat() if t.created_at else None,
                    "expires_at": t.expires_at.isoformat() if t.expires_at else None,
                    "token_hash": t.token_hash
                }
                for t in tokens
            ]
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Revoke an access token.")
async def admin_revoke_token(
    token_hash_prefix: Annotated[str, Field(description="Prefix of the token hash to revoke; at least 8 characters")]
) -> dict:
    """
    Révoque un token.
    
    Args:
        token_hash_prefix: Début du hash du token (8+ caractères)
        
    Returns:
        Statut de la révocation
    """
    try:
        # Sécurité : permission admin requise
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err
        
        # Trouver le token par son préfixe
        tokens = await get_tokens().list_tokens(include_revoked=False)
        
        matching = [t for t in tokens if t.token_hash.startswith(token_hash_prefix)]
        
        if not matching:
            return {"status": "error", "message": "Token not found"}
        
        if len(matching) > 1:
            return {"status": "error", "message": "Ambiguous prefix; provide more characters"}
        
        # Révoquer
        success = await get_tokens().revoke_token(matching[0].token_hash)
        
        if success:
            return {
                "status": "ok",
                "message": f"Token revoked for '{matching[0].client_name}'"
            }
        return {"status": "error", "message": "Revocation failed"}
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Update an access token.")
async def admin_update_token(
    token_hash_prefix: Annotated[str, Field(description="Token hash prefix; at least 8 characters")],
    add_memories: Annotated[Optional[List[str]], Field(default=None, description="Memories to add (for example, ['LEGAL', 'CLOUD'])")] = None,
    remove_memories: Annotated[Optional[List[str]], Field(default=None, description="Memories to remove (for example, ['LEGAL'])")] = None,
    set_memories: Annotated[Optional[List[str]], Field(default=None, description="Remplacer la liste (ex: ['CLOUD'], ou [] pour tout autoriser)")] = None,
    set_permissions: Annotated[Optional[List[str]], Field(default=None, description="Remplacer les permissions (ex: ['admin', 'read', 'write'] pour promouvoir en admin)")] = None,
    set_email: Annotated[Optional[str], Field(default=None, description="Modifier l'adresse email du token")] = None
) -> dict:
    """
    Met à jour les mémoires autorisées, les permissions et/ou l'email d'un token.
    
    Gestion des mémoires (mutuellement exclusifs avec set_memories) :
    - add_memories: Ajoute des mémoires à la liste existante
    - remove_memories: Retire des mémoires de la liste existante
    - set_memories: Remplace toute la liste ([] = accès à TOUTES les mémoires)
    
    Gestion des permissions :
    - set_permissions: Remplace les permissions (ex: ["admin", "read", "write"] pour promouvoir en admin)
    
    Permissions valides : 'read', 'write', 'admin'.
    Un token avec la permission 'admin' a accès à TOUS les outils, y compris la gestion des tokens.
    
    Args:
        token_hash_prefix: Début du hash du token (8+ caractères)
        add_memories: Mémoires à ajouter (ex: ["JURIDIQUE", "CLOUD"])
        remove_memories: Mémoires à retirer (ex: ["JURIDIQUE"])
        set_memories: Remplacer toute la liste (ex: ["CLOUD"], ou [] pour tout autoriser)
        set_permissions: Remplacer les permissions (ex: ["admin", "read", "write"])
        
    Returns:
        Anciennes et nouvelles mémoires/permissions autorisées
    """
    try:
        # Sécurité : permission admin requise
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err
        
        # Trouver le token par son préfixe
        tokens = await get_tokens().list_tokens(include_revoked=False)
        matching = [t for t in tokens if t.token_hash.startswith(token_hash_prefix)]
        
        if not matching:
            return {"status": "error", "message": "Token not found"}
        
        if len(matching) > 1:
            return {"status": "error", "message": "Ambiguous prefix; provide more characters"}
        
        result_parts = {}
        
        # === Mise à jour des permissions (si demandé) ===
        if set_permissions is not None:
            valid_perms = {"read", "write", "admin"}
            invalid = set(set_permissions) - valid_perms
            if invalid:
                return {
                    "status": "error",
                    "message": f"Permissions invalides: {invalid}. Valides: {sorted(valid_perms)}"
                }
            
            perm_result = await get_tokens().update_token_permissions(
                token_hash=matching[0].token_hash,
                permissions=set_permissions
            )
            if perm_result:
                result_parts["previous_permissions"] = perm_result["previous_permissions"]
                result_parts["current_permissions"] = perm_result["current_permissions"]
        
        # === Mise à jour des mémoires (si demandé) ===
        has_memory_update = add_memories or remove_memories or set_memories is not None
        if has_memory_update:
            # Vérifier que les mémoires existent (si on en ajoute)
            memories_to_check = (add_memories or []) + (set_memories or []) if set_memories is not None else (add_memories or [])
            if memories_to_check:
                existing_memories = await get_graph().list_memories()
                existing_ids = {m.id for m in existing_memories}
                unknown = [m for m in memories_to_check if m not in existing_ids]
                if unknown:
                    return {
                        "status": "error",
                        "message": f"Unknown memories: {unknown}. Available: {sorted(existing_ids)}"
                    }
            
            mem_result = await get_tokens().update_token_memories(
                token_hash=matching[0].token_hash,
                add_memories=add_memories,
                remove_memories=remove_memories,
                set_memories=set_memories
            )
            if mem_result:
                result_parts["previous_memories"] = mem_result["previous_memories"]
                result_parts["current_memories"] = mem_result["current_memories"]
        
        # === Mise à jour de l'email (si demandé) ===
        if set_email is not None:
            email_result = await get_tokens().update_token_email(
                token_hash=matching[0].token_hash,
                email=set_email
            )
            if email_result:
                result_parts["previous_email"] = email_result["previous_email"]
                result_parts["current_email"] = email_result["current_email"]
        
        if not result_parts:
            return {"status": "error", "message": "No changes requested (specify set_permissions, set_email, add_memories, remove_memories, or set_memories)"}
        
        # Construire le message
        messages = []
        if "current_permissions" in result_parts:
            perms = result_parts["current_permissions"]
            if "admin" in perms:
                messages.append(f"🔑 Promu ADMIN (permissions: {perms})")
            else:
                messages.append(f"Permissions: {perms}")
        if "current_memories" in result_parts:
            mems = result_parts["current_memories"]
            if not mems:
                messages.append("Access to all memories")
            else:
                messages.append(f"Memories: {mems}")
        
        return {
            "status": "ok",
            "client_name": matching[0].client_name,
            "token_hash_prefix": matching[0].token_hash[:8] + "...",
            **result_parts,
            "message": " | ".join(messages)
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# OUTILS MCP - Diagnostic
# =============================================================================

@mcp.tool(description="Return a memory's complete graph.")
async def memory_graph(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    format: Annotated[str, Field(default="full", description="Format : 'full' (tout), 'nodes', 'edges', 'documents'")] = "full"
) -> dict:
    """
    Récupère le graphe complet d'une mémoire (entités, relations et documents).
    
    Utile pour visualiser ou exporter le graphe de connaissances.
    Inclut les documents avec leur URI S3 pour permettre la récupération.
    
    Args:
        memory_id: ID de la mémoire
        format: "full" (tout), "nodes" (entités+docs), "edges" (relations), "documents" (liste docs avec URI S3)
        
    Returns:
        nodes: Liste des entités et documents avec leurs propriétés
        edges: Liste des relations entre entités et documents
        documents: Liste des documents avec id, filename, uri S3
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        graph_data = await get_graph().get_full_graph(memory_id)
        
        if format == "nodes":
            return {
                "status": "ok",
                "memory_id": memory_id,
                "node_count": len(graph_data["nodes"]),
                "nodes": graph_data["nodes"]
            }
        elif format == "edges":
            return {
                "status": "ok",
                "memory_id": memory_id,
                "edge_count": len(graph_data["edges"]),
                "edges": graph_data["edges"]
            }
        elif format == "documents":
            return {
                "status": "ok",
                "memory_id": memory_id,
                "document_count": len(graph_data["documents"]),
                "documents": graph_data["documents"]
            }
        else:  # full
            return {
                "status": "ok",
                "memory_id": memory_id,
                "node_count": len(graph_data["nodes"]),
                "edge_count": len(graph_data["edges"]),
                "document_count": len(graph_data["documents"]),
                "nodes": graph_data["nodes"],
                "edges": graph_data["edges"],
                "documents": graph_data["documents"]
            }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="List the documents in a memory.")
async def document_list(
    memory_id: Annotated[str, Field(description="Memory identifier")]
) -> dict:
    """
    Liste tous les documents d'une mémoire.
    
    Args:
        memory_id: ID de la mémoire
        
    Returns:
        Liste des documents avec leurs métadonnées
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        graph_data = await get_graph().get_full_graph(memory_id)
        docs = graph_data.get("documents", [])
        
        return {
            "status": "ok",
            "memory_id": memory_id,
            "count": len(docs),
            "documents": docs
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Return a document and optional decoded content.")
async def document_get(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    document_id: Annotated[str, Field(description="ID du document (UUID)")],
    include_content: Annotated[bool, Field(default=False, description="Download and include S3 content; this can be slow")] = False,
    content_format: Annotated[str, Field(default="text", description="Binary-file content format: 'text' for extracted text (default) or 'raw' for original bytes encoded as base64")] = "text"
) -> dict:
    """
    Récupère les métadonnées d'un document, et optionnellement son contenu.
    
    Par défaut, retourne uniquement les métadonnées (rapide, pas de téléchargement S3).
    Passez include_content=True pour télécharger et inclure le contenu du document.
    
    Pour les fichiers binaires (DOCX, PDF, XLSX...), deux modes :
    - content_format="text" (défaut) : extrait le texte lisible (paragraphes, tableaux)
    - content_format="raw" : retourne les bytes originaux en base64 (pour forwarding vers d'autres services)
    
    Args:
        memory_id: ID de la mémoire
        document_id: ID du document
        include_content: Si True, télécharge et inclut le contenu S3 (lent). Défaut: False.
        content_format: "text" (texte extrait, défaut) ou "raw" (base64 original). Défaut: "text".
        
    Returns:
        Métadonnées du document (et contenu si demandé)
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        
        # Récupérer les infos du document depuis le graphe (rapide, pas de S3)
        doc_info = await get_graph().get_document(memory_id, document_id)
        
        if not doc_info:
            return {"status": "error", "message": f"Document '{document_id}' not found"}
        
        result = {
            "status": "ok",
            "document": {
                "id": doc_info.get("id"),
                "filename": doc_info.get("filename"),
                "uri": doc_info.get("uri"),
                "hash": doc_info.get("hash"),
                "sha256": doc_info.get("sha256"),
                "ingested_at": doc_info.get("ingested_at"),
                "source_path": doc_info.get("source_path"),
                "repo_path": doc_info.get("repo_path"),
                "source_modified_at": doc_info.get("source_modified_at"),
                "size_bytes": doc_info.get("size_bytes", 0),
                "text_length": doc_info.get("text_length", 0),
                "content_type": doc_info.get("content_type"),
                "ingestion_status": doc_info.get("ingestion_status", "unknown"),
                "last_ingest_job_id": doc_info.get("last_ingest_job_id"),
                "chunk_count": doc_info.get("chunk_count", 0),
            },
        }
        
        # Télécharger le contenu S3 seulement si demandé
        if include_content and doc_info.get("uri"):
            try:
                uri = doc_info["uri"]
                content_bytes = await get_storage().download_document(memory_id, uri)

                # Distinguer fichiers texte vs binaires
                content_type = doc_info.get("content_type", "")
                text_extensions = {"txt", "md", "csv", "html", "htm", "json", "xml", "yaml", "yml"}

                if content_type in text_extensions:
                    # Fichiers texte : décodage UTF-8 direct (content_format ignoré)
                    result["content"] = content_bytes.decode('utf-8', errors='replace')
                elif content_format == "raw":
                    # Mode raw : bytes bruts en base64 (pour forwarding vers MCP Office etc.)
                    result["content_base64"] = base64.b64encode(content_bytes).decode('ascii')
                    result["content_format"] = "raw"
                    result["content_note"] = f"Original binary file ({content_type}) encoded as base64"
                else:
                    # Mode text (défaut) : extraction texte lisible
                    filename = doc_info.get("filename", "document")
                    extracted_text = _extract_text(content_bytes, filename)
                    if extracted_text:
                        result["content"] = extracted_text
                        result["content_format"] = "text"
                        result["content_note"] = f"Texte extrait du document binaire ({content_type}). Passez content_format='raw' pour le fichier original en base64."
                    else:
                        # Extraction impossible → fallback automatique vers raw
                        result["content_base64"] = base64.b64encode(content_bytes).decode('ascii')
                        result["content_format"] = "raw"
                        result["content_note"] = f"Text extraction is unsupported for {content_type}; using automatic base64 fallback."
            except Exception as e:
                result["content"] = f"[Erreur lecture S3: {e}]"
        
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Delete a document and its graph/vector data.")
async def document_delete(
    memory_id: Annotated[str, Field(description="Memory identifier")],
    document_id: Annotated[str, Field(description="Identifier of the document to delete (UUID)")]
) -> dict:
    """
    Supprime un document du graphe ET de S3.
    
    Supprime :
    - Le fichier S3 associé
    - Le nœud Document dans Neo4j
    - Les relations MENTIONS du document
    - Les entités orphelines (non mentionnées par d'autres documents)
    - Les relations RELATED_TO impliquant des entités orphelines
    
    Args:
        memory_id: ID de la mémoire
        document_id: ID du document à supprimer
        
    Returns:
        Statut de la suppression avec compteurs (graphe + S3)
    """
    try:
        # Vérifier l'accès à la mémoire
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err

        # P7-8 : suppression = mutation destructive → permission write requise
        # (l'accès en lecture seule ne suffit pas). Aligne document_delete sur
        # les autres mutations (memory_ingest, memory_update, ...).
        write_err = check_write_permission()
        if write_err:
            return write_err

        # Suppression multi-backend ordonnée (Qdrant → Neo4j → S3), compensable
        from .core.ingest_pipeline import delete_document_everywhere
        result = await delete_document_everywhere(memory_id, document_id)
        errors = result.get("errors", [])

        if result.get("neo4j_deleted"):
            # Si un backend a échoué (Qdrant/S3), le signaler explicitement à l'appelant
            return {
                "status": "partial_deleted" if errors else "deleted",
                "document_id": document_id,
                "relations_deleted": result.get("relations_deleted", 0),
                "entities_deleted": result.get("entities_deleted", 0),
                "qdrant_chunks_deleted": result.get("qdrant_chunks_deleted", 0),
                "s3_deleted": result.get("s3_deleted", False),
                "errors": errors,
                "message": ("Document deleted, but cleanup is incomplete (see errors)"
                            if errors else "Document deleted"),
            }
        return {"status": "error", "message": "Document not found or Neo4j deletion failed", "errors": errors}

    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="List available ontologies.")
async def ontology_list() -> dict:
    """
    Liste toutes les ontologies disponibles.
    
    Les ontologies définissent les règles d'extraction pour différents domaines.
    Chaque mémoire DOIT avoir une ontologie. Exemples:
    - legal: Documents juridiques et contractuels
    - cloud: Infrastructure cloud et certifications
    - managed-services: Infogérance et services managés
    - technical: Documentation technique et API
    
    Returns:
        Liste des ontologies avec leurs métadonnées
    """
    try:
        from .core.ontology import get_ontology_manager
        ontology_manager = get_ontology_manager()
        ontologies = ontology_manager.list_ontologies()
        
        return {
            "status": "ok",
            "count": len(ontologies),
            "ontologies": ontologies
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


def _safe_ontology_name(name: str) -> str:
    """Valide un nom d'ontologie pour éviter path traversal et noms ambigus."""
    import re
    value = (name or "").strip()
    if not value:
        raise ValueError("Ontology name is required")
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_-]{0,79}", value):
        raise ValueError("Invalid ontology name (letters, digits, _, and - only)")
    return value


def _ontology_dir() -> str:
    """Retourne le dossier ONTOLOGIES actif."""
    from .core.ontology import get_ontology_manager
    ontology_manager = get_ontology_manager()
    path = getattr(ontology_manager, "_ontology_path", None)
    if not path or not os.path.isdir(path):
        raise ValueError("ONTOLOGIES directory not found")
    return path


def _ontology_file_for_name(name: str) -> Optional[str]:
    """Trouve le fichier YAML correspondant au champ name d'une ontologie."""
    safe_name = _safe_ontology_name(name)
    directory = _ontology_dir()
    direct = os.path.join(directory, f"{safe_name}.yaml")
    if os.path.exists(direct):
        return direct

    import yaml
    for filename in os.listdir(directory):
        if not filename.endswith((".yaml", ".yml")):
            continue
        filepath = os.path.join(directory, filename)
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f) or {}
            if data.get("name") == safe_name:
                return filepath
        except Exception:
            continue
    return None


def _load_ontology_yaml(content_yaml: str) -> dict:
    """Parse et valide le minimum structurel d'une ontologie YAML."""
    import yaml
    try:
        data = yaml.safe_load(content_yaml) or {}
    except yaml.YAMLError as e:
        raise ValueError(f"Invalid YAML: {e}")

    name = _safe_ontology_name(str(data.get("name", "")))
    if not isinstance(data.get("entity_types"), list) or not data["entity_types"]:
        raise ValueError("entity_types must be a non-empty list")
    if not isinstance(data.get("relation_types"), list) or not data["relation_types"]:
        raise ValueError("relation_types must be a non-empty list")
    data["name"] = name
    return data


@mcp.tool(description="Return an ontology's parsed definition.")
async def ontology_get(
    name: Annotated[str, Field(description="Ontology name")]
) -> dict:
    """
    Lit une ontologie avec son contenu YAML brut.

    Args:
        name: Nom de l'ontologie

    Returns:
        Métadonnées + contenu YAML
    """
    try:
        _safe_ontology_name(name)
        filepath = _ontology_file_for_name(name)
        if not filepath:
            return {"status": "error", "message": f"Ontology '{name}' not found"}

        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        data = _load_ontology_yaml(content)
        return {
            "status": "ok",
            "name": data.get("name"),
            "version": str(data.get("version", "")),
            "description": data.get("description", ""),
            "filename": os.path.basename(filepath),
            "entity_types_count": len(data.get("entity_types", [])),
            "relation_types_count": len(data.get("relation_types", [])),
            "content": content,
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Export an ontology as YAML.")
async def ontology_export(
    name: Annotated[str, Field(description="Ontology name")]
) -> dict:
    """
    Exporte une ontologie en YAML et base64.

    Args:
        name: Nom de l'ontologie

    Returns:
        Contenu YAML + content_base64 pour téléchargement
    """
    result = await ontology_get(name)
    if result.get("status") != "ok":
        return result
    content = result.get("content", "")
    return {
        "status": "ok",
        "name": result.get("name"),
        "filename": result.get("filename") or f"{name}.yaml",
        "content": content,
        "content_base64": base64.b64encode(content.encode("utf-8")).decode("ascii"),
    }


@mcp.tool(description="Import an ontology from YAML.")
async def ontology_import(
    content_yaml: Annotated[str, Field(description="Ontology YAML content")],
    overwrite: Annotated[bool, Field(default=False, description="Overwrite an existing ontology")] = False
) -> dict:
    """
    Importe une nouvelle ontologie dans le référentiel ONTOLOGIES/.

    Requiert la permission admin.
    """
    try:
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err

        data = _load_ontology_yaml(content_yaml)
        name = data["name"]
        directory = _ontology_dir()
        existing = _ontology_file_for_name(name)
        if existing and not overwrite:
            return {
                "status": "error",
                "message": f"Ontology '{name}' already exists. Use overwrite=true to replace it."
            }

        filepath = existing or os.path.join(directory, f"{name}.yaml")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content_yaml.rstrip() + "\n")

        from .core.ontology import get_ontology_manager
        get_ontology_manager().reload()
        return {
            "status": "ok",
            "name": name,
            "filename": os.path.basename(filepath),
            "message": f"Ontology '{name}' imported",
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Update an existing ontology from YAML.")
async def ontology_update(
    name: Annotated[str, Field(description="Name of the ontology to update")],
    content_yaml: Annotated[str, Field(description="Nouveau contenu YAML complet")]
) -> dict:
    """
    Remplace le contenu YAML d'une ontologie existante.

    Requiert la permission admin.
    """
    try:
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err

        safe_name = _safe_ontology_name(name)
        filepath = _ontology_file_for_name(safe_name)
        if not filepath:
            return {"status": "error", "message": f"Ontology '{safe_name}' not found"}

        data = _load_ontology_yaml(content_yaml)
        if data["name"] != safe_name:
            return {
                "status": "error",
                "message": "The YAML 'name' field must remain unchanged. Use import to create a new ontology."
            }

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content_yaml.rstrip() + "\n")

        from .core.ontology import get_ontology_manager
        get_ontology_manager().reload()
        return {
            "status": "ok",
            "name": safe_name,
            "filename": os.path.basename(filepath),
            "message": f"Ontology '{safe_name}' updated",
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Delete an ontology.")
async def ontology_delete(
    name: Annotated[str, Field(description="Name of the ontology to delete")],
    force: Annotated[bool, Field(default=False, description="Delete even when memories use the ontology")] = False
) -> dict:
    """
    Supprime une ontologie du référentiel ONTOLOGIES/.

    Par défaut, refuse la suppression si une mémoire existante utilise cette
    ontologie. Requiert la permission admin.
    """
    try:
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err

        safe_name = _safe_ontology_name(name)
        filepath = _ontology_file_for_name(safe_name)
        if not filepath:
            return {"status": "error", "message": f"Ontology '{safe_name}' not found"}

        memories = await get_graph().list_memories()
        users = [m.id for m in memories if m.ontology == safe_name]
        if users and not force:
            return {
                "status": "error",
                "message": f"Ontology is used by {len(users)} memory or memories: {users}. Use force=true to delete it intentionally.",
                "used_by": users,
            }

        os.remove(filepath)
        from .core.ontology import get_ontology_manager
        get_ontology_manager().reload()
        return {
            "status": "deleted",
            "name": safe_name,
            "filename": os.path.basename(filepath),
            "used_by": users,
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Check graph, vector, and object-storage consistency.")
async def storage_check(
    memory_id: Annotated[Optional[str], Field(default=None, description="Optional memory identifier; check all memories when omitted")] = None
) -> dict:
    """
    Vérifie la cohérence entre le graphe Neo4j et le stockage S3.
    
    Pour chaque mémoire (ou une mémoire spécifique) :
    1. Vérifie que chaque document du graphe est accessible sur S3
    2. Détecte les fichiers orphelins sur S3 (pas de référence dans le graphe)
    3. Retourne un rapport complet avec statistiques
    
    Args:
        memory_id: ID d'une mémoire spécifique (optionnel, toutes si omis)
        
    Returns:
        Rapport de cohérence S3/Graphe avec documents OK, manquants et orphelins
    """
    try:
        # Sécurité : admin requis pour le mode global, check_memory_access pour une mémoire spécifique
        if memory_id:
            access_err = check_memory_access(memory_id)
            if access_err:
                return access_err
        else:
            admin_err = check_admin_permission()
            if admin_err:
                return admin_err
        
        # 1. Récupérer les mémoires à vérifier
        if memory_id:
            memory = await get_graph().get_memory(memory_id)
            if not memory:
                return {"status": "error", "message": f"Memory '{memory_id}' not found"}
            memories = [memory]
        else:
            memories = await get_graph().list_memories()
        
        # 2. Collecter toutes les URIs des documents référencés dans le graphe
        graph_uris = set()          # URIs référencées dans Neo4j
        graph_uri_details = {}      # URI -> {memory_id, filename, doc_id}
        memory_prefixes = set()     # Préfixes S3 des mémoires connues
        
        for mem in memories:
            mid = mem.id
            memory_prefixes.add(f"{mid}/")
            graph_data = await get_graph().get_full_graph(mid)
            
            for doc in graph_data.get("documents", []):
                uri = doc.get("uri", "")
                if uri:
                    graph_uris.add(uri)
                    graph_uri_details[uri] = {
                        "memory_id": mid,
                        "filename": doc.get("filename", "?"),
                        "doc_id": doc.get("id", "?")
                    }
        
        # 3. Vérifier l'accessibilité S3 de chaque document du graphe
        check_result = await get_storage().check_documents(list(graph_uris))
        
        # Enrichir les détails avec les infos du graphe
        for detail in check_result.get("details", []):
            uri = detail.get("uri", "")
            if uri in graph_uri_details:
                detail["memory_id"] = graph_uri_details[uri]["memory_id"]
                detail["filename"] = graph_uri_details[uri]["filename"]
                detail["doc_id"] = graph_uri_details[uri]["doc_id"]
        
        # 4. Lister tous les objets S3 pour détecter les orphelins
        #    IMPORTANT : pour la détection d'orphelins, on compare avec TOUTES
        #    les mémoires, pas seulement celles du scope. Sinon les docs des
        #    autres mémoires apparaissent comme faux-positifs.
        all_s3_objects = await get_storage().list_all_objects()
        
        # Collecter les clés S3 de TOUTES les mémoires (pas seulement le scope)
        all_graph_uris = set(graph_uris)  # Commencer avec celles du scope
        if memory_id:
            # Charger les URIs des autres mémoires aussi
            all_memories = await get_graph().list_memories()
            for mem in all_memories:
                if mem.id == memory_id:
                    continue  # Déjà chargé
                other_graph = await get_graph().get_full_graph(mem.id)
                for doc in other_graph.get("documents", []):
                    uri = doc.get("uri", "")
                    if uri:
                        all_graph_uris.add(uri)
        
        # Convertir les URIs du graphe en clés S3 pour comparaison
        graph_keys = set()
        for uri in all_graph_uris:
            try:
                key = get_storage()._parse_key(uri)
                graph_keys.add(key)
            except ValueError:
                pass
        
        # Ajouter les ontologies comme fichiers légitimes (pas orphelins)
        # Les fichiers _ontology_*.yaml sont des fichiers de config, pas des orphelins
        
        # Détecter les orphelins : sur S3 mais pas dans le graphe
        orphans = []
        for obj in all_s3_objects:
            key = obj["key"]
            
            # Ignorer les fichiers de health check
            if key.startswith("_health_check/"):
                continue
            
            # Ignorer les backups (gérés séparément via backup_list)
            if key.startswith("_backups/"):
                continue
            
            # Ignorer les ontologies (fichiers légitimes)
            # Le pattern est {hash[:8]}__ontology_{name}.yaml (double _ car hash + _ontology)
            if "_ontology_" in key:
                continue
            
            # Si la clé n'est pas référencée dans le graphe → orphelin
            if key not in graph_keys:
                orphans.append({
                    "key": key,
                    "uri": obj["uri"],
                    "size": obj["size"],
                    "last_modified": obj["last_modified"]
                })
        
        # 5. Construire le rapport
        def _human_size(size_bytes):
            """Convertit des bytes en taille lisible."""
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024:
                    return f"{size_bytes:.1f} {unit}"
                size_bytes /= 1024
            return f"{size_bytes:.1f} TB"
        
        orphan_total_size = sum(o["size"] for o in orphans)

        # 4b. Cohérence multi-backend (Qdrant + doublons source_path + ingestions partielles)
        duplicate_source_paths = []   # [{memory_id, source_path, doc_ids}]
        partial_ingestions = []       # [{memory_id, doc_id, source_path, ingestion_status}]
        qdrant_orphan_vectors = []    # [{memory_id, doc_id}] : chunks Qdrant sans Document Neo4j
        qdrant_missing_chunks = []    # [{memory_id, doc_id, source_path}] : doc succeeded mais chunks ≠
        qdrant_errors = []            # [{memory_id, error}] : check Qdrant incomplet (panne transitoire)
        for mem in memories:
            mid = mem.id
            mem_graph = await get_graph().get_full_graph(mid)
            docs = mem_graph.get("documents", [])
            neo4j_doc_ids = {d.get("id") for d in docs if d.get("id")}

            # Doublons source_path
            by_sp = {}
            for d in docs:
                sp = d.get("source_path")
                if sp:
                    by_sp.setdefault(sp, []).append(d.get("id"))
            for sp, ids in by_sp.items():
                if len(ids) > 1:
                    duplicate_source_paths.append({"memory_id": mid, "source_path": sp, "doc_ids": ids})

            # Ingestions partielles (statut durable ni succeeded ni legacy)
            for d in docs:
                st = d.get("ingestion_status")
                if st in ("running", "failed", "cleanup_pending"):
                    partial_ingestions.append({
                        "memory_id": mid, "doc_id": d.get("id"),
                        "source_path": d.get("source_path"), "ingestion_status": st,
                    })

            # Qdrant : vecteurs orphelins + chunks manquants/partiels
            try:
                qdrant_doc_ids = await get_vector_store().list_doc_ids(mid)
                for did in qdrant_doc_ids - neo4j_doc_ids:
                    qdrant_orphan_vectors.append({"memory_id": mid, "doc_id": did})
                for d in docs:
                    expected = d.get("chunk_count") or 0
                    if d.get("ingestion_status") == "succeeded" and expected > 0:
                        # Comparer le compte RÉEL de vecteurs Qdrant à chunk_count (détecte le partiel)
                        actual = await get_vector_store().count_document_chunks(mid, d.get("id"))
                        if actual != expected:
                            qdrant_missing_chunks.append({
                                "memory_id": mid, "doc_id": d.get("id"),
                                "source_path": d.get("source_path"),
                                "expected_chunks": expected, "actual_chunks": actual,
                            })
            except Exception as e:
                # Ne PAS prétendre "0 incohérence" : on signale que le check Qdrant
                # de cette mémoire est incomplet (le client doit en tenir compte).
                print(f"⚠️ [storage_check] Incomplete Qdrant data for {mid}: {e}", file=sys.stderr)
                qdrant_errors.append({"memory_id": mid, "error": str(e)})

        consistency_issues = (
            len(duplicate_source_paths) + len(partial_ingestions)
            + len(qdrant_orphan_vectors) + len(qdrant_missing_chunks)
        )

        report = {
            "status": "ok",
            "scope": memory_id or "all",
            "memories_checked": len(memories),
            "consistency": {
                "issues": consistency_issues,
                "complete": len(qdrant_errors) == 0,
                "duplicate_source_paths": duplicate_source_paths,
                "partial_ingestions": partial_ingestions,
                "qdrant_orphan_vectors": qdrant_orphan_vectors,
                "qdrant_missing_chunks": qdrant_missing_chunks,
                "qdrant_errors": qdrant_errors,
            },
            "graph_documents": {
                "total": check_result["total"],
                "accessible": check_result["accessible"],
                "missing": check_result["missing"],
                "errors": check_result["errors"],
                "total_size": _human_size(check_result["total_size_bytes"]),
                "total_size_bytes": check_result["total_size_bytes"],
                "details": check_result["details"]
            },
            "s3_orphans": {
                "count": len(orphans),
                "total_size": _human_size(orphan_total_size),
                "total_size_bytes": orphan_total_size,
                "files": orphans
            },
            "s3_total_objects": len(all_s3_objects),
            "summary": (
                f"✅ {check_result['accessible']}/{check_result['total']} docs accessibles"
                + (f", ❌ {check_result['missing']} manquants" if check_result['missing'] > 0 else "")
                + (f", ⚠️ {len(orphans)} orphelins S3 ({_human_size(orphan_total_size)})" if orphans else "")
                + (f", 🔗 {consistency_issues} incohérence(s) (Qdrant/source_path/partiels)" if consistency_issues else "")
                + (f", ⚠️ check Qdrant incomplet sur {len(qdrant_errors)} mémoire(s)" if qdrant_errors else "")
            )
        }
        
        return report
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Find or delete orphaned storage objects.")
async def storage_cleanup(
    dry_run: Annotated[bool, Field(default=True, description="List only when true; permanently delete orphans when false")] = True
) -> dict:
    """
    Nettoie les fichiers orphelins sur S3.
    
    Un fichier orphelin est un objet S3 qui n'est référencé par aucun document
    dans le graphe Neo4j (ni par une ontologie de mémoire).
    
    ⚠️ Par défaut, mode dry_run=True : liste les fichiers sans les supprimer.
    Passez dry_run=False pour effectuer la suppression.
    
    Args:
        dry_run: Si True, liste seulement. Si False, supprime réellement.
        
    Returns:
        Liste des fichiers orphelins (supprimés ou à supprimer)
    """
    try:
        # Sécurité : permission admin requise (opération globale S3)
        admin_err = check_admin_permission()
        if admin_err:
            return admin_err
        
        # 1. Exécuter le check complet pour identifier les orphelins
        check = await storage_check()
        
        if check.get("status") != "ok":
            return check
        
        orphans = check.get("s3_orphans", {}).get("files", [])
        
        if not orphans:
            return {
                "status": "ok",
                "message": "No orphaned files found. S3 is clean.",
                "orphans_found": 0,
                "deleted": 0,
                "dry_run": dry_run
            }
        
        if dry_run:
            return {
                "status": "ok",
                "message": f"Found {len(orphans)} orphaned files ({check['s3_orphans']['total_size']}). "
                           f"Relancez avec dry_run=false pour les supprimer.",
                "orphans_found": len(orphans),
                "deleted": 0,
                "dry_run": True,
                "files": orphans
            }
        
        # 2. Supprimer les orphelins
        keys_to_delete = [o["key"] for o in orphans]
        delete_result = await get_storage().delete_objects(keys_to_delete)
        
        return {
            "status": "ok",
            "message": f"Deleted {delete_result['deleted_count']} orphaned files "
                       f"({check['s3_orphans']['total_size']} libérés).",
            "orphans_found": len(orphans),
            "deleted": delete_result["deleted_count"],
            "errors": delete_result["error_count"],
            "dry_run": False,
            "files": orphans
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Return public service identity and capability information.")
async def system_about() -> dict:
    """
    Identity and role of Hivemind's embedded long-memory runtime.
    
    Retourne une description complète du service :
    - Qui il est et pourquoi il existe
    - Ses capacités (outils disponibles, ontologies)
    - Les mémoires actives
    - Les services connectés et leur état
    
    Retourne l'identité et les capacités publiques sans authentification.
    Les détails sensibles (mémoires, services, configuration) nécessitent
    une permission 'read' minimum (sécurité v2.1.0).
    
    Returns:
        Identité, capacités, et si authentifié: mémoires actives, état des services
    """
    try:
        # Version depuis le fichier VERSION
        version = "?"
        try:
            version_path = os.path.join(os.path.dirname(__file__), "..", "..", "VERSION")
            if os.path.exists(version_path):
                with open(version_path) as f:
                    version = f.read().strip()
        except Exception:
            pass
        
        # Sensitive service/configuration details are returned only when the
        # caller has an authenticated Hivemind identity.
        auth = current_auth.get()
        is_authenticated = auth is not None
        if auth and auth.get("type") == "token":
            permissions = auth.get("permissions", [])
            is_authenticated = any(
                permission in permissions
                for permission in ("read", "write", "manage", "admin")
            )
        
        # Ontologies disponibles (public)
        ontologies_info = []
        try:
            from .core.ontology import get_ontology_manager
            ontology_manager = get_ontology_manager()
            ontologies_info = ontology_manager.list_ontologies()
        except Exception:
            pass
        
        # Mémoires actives (authentifié uniquement)
        memories_info = []
        if is_authenticated:
            try:
                memories = await get_graph().list_memories()
                for m in memories:
                    stats = await get_graph().get_memory_stats(m.id)
                    memories_info.append({
                        "id": m.id,
                        "name": m.name,
                        "ontology": m.ontology,
                        "documents": stats.document_count,
                        "entities": stats.entity_count,
                        "relations": stats.relation_count,
                    })
            except Exception:
                pass
        
        # État des services (authentifié uniquement)
        services_status = {}
        if is_authenticated:
            for name, test_fn in [
                ("neo4j", lambda: get_graph().test_connection()),
                ("s3", lambda: get_storage().test_connection()),
                ("qdrant", lambda: get_vector_store().test_connection()),
                ("llmaas", lambda: get_extractor().test_connection()),
                ("embedding", lambda: get_embedder().test_connection()),
            ]:
                try:
                    result = await test_fn()
                    services_status[name] = result.get("status", "unknown")
                except Exception:
                    services_status[name] = "error"
        
        # Outils MCP disponibles (comptage par catégorie)
        tools_categories = {
            "Gestion mémoires": ["memory_create", "memory_update", "memory_delete", "memory_list", "memory_stats"],
            "Ingestion": ["memory_ingest"],
            "Ingestion asynchrone": ["memory_ingest_async", "memory_ingest_batch_async", "ingest_job_status", "ingest_job_list", "ingest_job_cancel"],
            "Recherche & Q&A": ["memory_search", "memory_query", "memory_get_context", "question_answer"],
            "Documents": ["document_list", "document_get", "document_delete"],
            "Ontologies": ["ontology_list", "ontology_get", "ontology_export", "ontology_import", "ontology_update", "ontology_delete"],
            "Backup/Restore": ["backup_create", "backup_list", "backup_restore", "backup_download", "backup_delete", "backup_restore_archive"],
            "Administration": ["admin_create_token", "admin_list_tokens", "admin_revoke_token", "admin_update_token"],
            "Diagnostic": ["system_health", "system_about", "system_whoami", "storage_check", "storage_cleanup"],
            "Visualisation": ["memory_graph"],
        }
        total_tools = sum(len(v) for v in tools_categories.values())
        
        return {
            "status": "ok",
            "identity": {
                "name": settings.mcp_server_name,
                "version": version,
                "description": "Embedded ontology and knowledge-graph runtime "
                               "for Hivemind's derived long-memory tier.",
                "purpose": "Derive searchable entities, relations, and semantic "
                           "context from documents within a Hivemind space.",
                "approach": "Neo4j and Qdrant form a derived projection only. "
                            "This runtime is never authoritative for Hivemind "
                            "commits, rollback, audit, membership, or recovery.",
                "provider": "Lesur AI",
                "repo": "https://github.com/Lesur-ai/hivemind",
                "upstream": "https://github.com/cloud-temple/graph-memory",
            },
            "capabilities": {
                "total_tools": total_tools,
                "categories": {k: len(v) for k, v in tools_categories.items()},
                "tools_list": tools_categories,
                "ontologies": [
                    {"name": o.get("name", "?"), "description": o.get("description", "")}
                    for o in ontologies_info
                ],
                "supported_formats": ["txt", "md", "html", "docx", "pdf", "csv"],
            },
            "memories": memories_info,
            "services": services_status,
            "configuration": (
                {
                    "llm_model": settings.llmaas_model,
                    "embedding_model": settings.llmaas_embedding_model,
                    "embedding_dimensions": settings.llmaas_embedding_dimensions,
                    "rag_score_threshold": settings.rag_score_threshold,
                    "chunk_size": settings.chunk_size,
                    "backup_retention": settings.backup_retention_count,
                }
                if is_authenticated
                else {}
            ),
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Return service and dependency health.")
async def system_health() -> dict:
    """
    Vérifie l'état de santé du système.
    
    Teste les connexions à tous les services :
    S3, Neo4j, LLMaaS, Qdrant, Embedding.
    Les 5 doivent être OK, sinon le service est considéré en erreur.
    
    Returns:
        État de chaque service
    """
    results = {}

    # P12-3 (Hivemind #268) : les échecs proxy peuvent embarquer l'URL proxy
    # brute (potentiellement porteuse de credentials) dans le message — chaque
    # message sortant de santé est donc redacted (no-op sans secret).

    # Test S3
    try:
        results["s3"] = await get_storage().test_connection()
    except Exception as e:
        results["s3"] = {"status": "error", "message": redact_proxy_secrets(str(e))}

    # Test Neo4j
    try:
        results["neo4j"] = await get_graph().test_connection()
    except Exception as e:
        results["neo4j"] = {"status": "error", "message": redact_proxy_secrets(str(e))}

    # Test LLMaaS (génération)
    try:
        results["llmaas"] = await get_extractor().test_connection()
    except Exception as e:
        results["llmaas"] = {"status": "error", "message": redact_proxy_secrets(str(e))}

    # Test Qdrant
    try:
        results["qdrant"] = await get_vector_store().test_connection()
    except Exception as e:
        results["qdrant"] = {"status": "error", "message": redact_proxy_secrets(str(e))}

    # Test Embedding (LLMaaS endpoint)
    try:
        results["embedding"] = await get_embedder().test_connection()
    except Exception as e:
        results["embedding"] = {"status": "error", "message": redact_proxy_secrets(str(e))}
    
    # Statut global : TOUS doivent être OK (couplage strict)
    all_ok = all(r.get("status") == "ok" for r in results.values())
    
    return {
        "status": "ok" if all_ok else "error",
        "services": results
    }


@mcp.tool(description="Return the current caller's identity and permissions.")
async def system_whoami() -> dict:
    """
    Identité du token courant.

    Retourne les informations du contexte d'authentification :
    - Type d'authentification (bootstrap, token, localhost)
    - Nom du client
    - Permissions (read, write, admin)
    - Mémoires autorisées
    - Email, date de création, expiration (si token)

    Utile pour vérifier son identité et ses droits d'accès.

    Returns:
        Identité et permissions du token courant
    """
    auth = current_auth.get()

    # Pas d'auth (localhost, endpoints publics)
    if auth is None:
        return {
            "status": "ok",
            "auth_type": "localhost",
            "client_name": "localhost",
            "permissions": ["admin", "read", "write"],
            "memory_ids": [],
            "note": "Local access without authentication"
        }

    result = {
        "status": "ok",
        "auth_type": auth.get("type", "unknown"),
        "client_name": auth.get("client_name", "unknown"),
        "permissions": auth.get("permissions", []),
        "memory_ids": auth.get("memory_ids", []),
    }

    # Pour les tokens, enrichir avec les infos stockées en base (email, dates)
    if auth.get("type") == "token" and auth.get("token_hash"):
        result["token_hash"] = auth["token_hash"]
        try:
            tokens = await get_tokens().list_tokens(include_revoked=False)
            for t in tokens:
                if t.token_hash == auth["token_hash"]:
                    result["email"] = t.email
                    result["created_at"] = t.created_at.isoformat() if t.created_at else None
                    result["expires_at"] = t.expires_at.isoformat() if t.expires_at else None
                    break
        except Exception:
            pass

    return result


# =============================================================================
# OUTILS MCP - Backup / Restore
# =============================================================================

@mcp.tool(description="Create a backup of one memory or all memories.")
async def backup_create(
    memory_id: Annotated[Optional[str], Field(default=None, description="Memory to back up; an admin backs up all memories when omitted")] = None,
    description: Annotated[Optional[str], Field(default=None, description="Optional backup description")] = None,
    ctx: Optional[Context] = None
) -> dict:
    """
    Crée un backup complet d'une mémoire, ou de toutes les mémoires pour un admin.
    
    Exporte le graphe Neo4j (entités, relations, documents),
    les vecteurs Qdrant (embeddings), et les références des documents S3.
    Applique la politique de rétention (BACKUP_RETENTION_COUNT).
    
    Args:
        memory_id: ID de la mémoire à sauvegarder. Si omis, backup de toutes les mémoires.
        description: Description optionnelle du backup
        
    Returns:
        backup_id, statistiques, temps d'exécution
    """
    try:
        memory_id = (memory_id or "").strip()

        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass

        if not memory_id:
            admin_err = check_admin_permission()
            if admin_err:
                return admin_err

            memories = await get_graph().list_memories()
            results = []
            errors = []
            for memory in memories:
                try:
                    result = await get_backup().create_backup(
                        memory_id=memory.id,
                        description=description,
                        progress_callback=_progress
                    )
                    results.append(result)
                except Exception as e:
                    errors.append({"memory_id": memory.id, "message": str(e)})

            return {
                "status": "ok" if not errors else "error",
                "mode": "all_memories",
                "requested_count": len(memories),
                "created_count": len(results),
                "error_count": len(errors),
                "backups": results,
                "errors": errors,
                "message": (
                    f"{len(results)} backup(s) created"
                    if not errors
                    else f"{len(results)} backup(s) created, {len(errors)} error(s)"
                ),
            }

        # Sécurité : vérifier accès mémoire + permission write
        access_err = check_memory_access(memory_id)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        result = await get_backup().create_backup(
            memory_id=memory_id,
            description=description,
            progress_callback=_progress
        )
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="List available memory backups.")
async def backup_list(
    memory_id: Annotated[Optional[str], Field(default=None, description="Optional memory filter; list all when omitted")] = None
) -> dict:
    """
    Liste les backups disponibles sur S3.
    
    Args:
        memory_id: Si fourni, liste uniquement les backups de cette mémoire.
                   Sinon, liste tous les backups.
        
    Returns:
        Liste des backups avec date, taille, statistiques
    """
    try:
        # Sécurité : filtrer par mémoires autorisées
        if memory_id:
            access_err = check_memory_access(memory_id)
            if access_err:
                return access_err
        
        backups = await get_backup().list_backups(memory_id=memory_id)
        
        # Filtrer les backups par les mémoires autorisées (si token restreint)
        allowed = get_allowed_memory_ids()
        # P7-4 (ADR-0019): fail-closed — no auth context => deny the listing.
        if allowed is DENY_ALL:
            return {"status": "error", "message": "Authentication required"}
        if allowed is not None and len(allowed) > 0 and not memory_id:
            backups = [b for b in backups if b.get("memory_id") in allowed]
        
        return {
            "status": "ok",
            "count": len(backups),
            "backups": [
                {
                    "backup_id": b.get("backup_id"),
                    "memory_id": b.get("memory_id"),
                    "memory_name": b.get("memory_name"),
                    "description": b.get("description"),
                    "created_at": b.get("created_at"),
                    "stats": b.get("stats", {}),
                    "version": b.get("version"),
                }
                for b in backups
            ]
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Restore a memory from an object-storage backup.")
async def backup_restore(
    backup_id: Annotated[str, Field(description="ID du backup (format: 'memory_id/timestamp')")],
    ctx: Optional[Context] = None
) -> dict:
    """
    Restaure une mémoire depuis un backup S3.
    
    ⚠️ La mémoire NE DOIT PAS exister (erreur sinon).
    Supprimez-la d'abord avec memory_delete si nécessaire.
    
    Restaure le graphe Neo4j + les vecteurs Qdrant tels qu'ils étaient,
    SANS refaire l'extraction LLM (instantané).
    
    Args:
        backup_id: ID du backup (format: "memory_id/timestamp")
        
    Returns:
        Compteurs de restauration (entités, relations, vecteurs, documents S3)
    """
    try:
        # Sécurité : extraire memory_id du backup_id, vérifier accès + write
        from .core.backup import BackupService
        mid, _ = BackupService._validate_backup_id(backup_id)
        access_err = check_memory_access(mid)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        result = await get_backup().restore_backup(
            backup_id=backup_id,
            progress_callback=_progress
        )
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Download a backup as a base64-encoded archive.")
async def backup_download(
    backup_id: Annotated[str, Field(description="ID du backup (format: 'memory_id/timestamp')")],
    include_documents: Annotated[bool, Field(default=False, description="Si true, inclut les documents originaux (PDF, DOCX, etc.)")] = False,
    ctx: Optional[Context] = None
) -> dict:
    """
    Télécharge un backup sous forme d'archive tar.gz encodée en base64.
    
    Par défaut (light) : uniquement les données JSON (graphe + vecteurs).
    Avec include_documents=True : inclut aussi les fichiers originaux (PDF, DOCX, etc.).
    
    Args:
        backup_id: ID du backup (format: "memory_id/timestamp")
        include_documents: Si True, inclut les documents originaux dans l'archive
        
    Returns:
        Archive tar.gz encodée en base64 + nom de fichier suggéré
    """
    try:
        # Sécurité : extraire memory_id du backup_id, vérifier accès mémoire
        from .core.backup import BackupService
        mid, _ = BackupService._validate_backup_id(backup_id)
        access_err = check_memory_access(mid)
        if access_err:
            return access_err
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        archive_bytes = await get_backup().download_backup(
            backup_id=backup_id,
            include_documents=include_documents,
            progress_callback=_progress
        )
        
        # Encoder en base64 pour transmission via MCP
        import base64 as b64
        archive_b64 = b64.b64encode(archive_bytes).decode("ascii")
        
        # Nom de fichier suggéré
        safe_id = backup_id.replace("/", "-")
        filename = f"backup-{safe_id}.tar.gz"
        
        return {
            "status": "ok",
            "backup_id": backup_id,
            "filename": filename,
            "size_bytes": len(archive_bytes),
            "include_documents": include_documents,
            "content_base64": archive_b64,
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Delete a memory backup.")
async def backup_delete(
    backup_id: Annotated[str, Field(description="Backup to delete in 'memory_id/timestamp' format")]
) -> dict:
    """
    Supprime un backup de S3.
    
    Args:
        backup_id: ID du backup (format: "memory_id/timestamp")
        
    Returns:
        Nombre de fichiers supprimés
    """
    try:
        # Sécurité : extraire memory_id du backup_id, vérifier accès + write
        from .core.backup import BackupService
        mid, _ = BackupService._validate_backup_id(backup_id)
        access_err = check_memory_access(mid)
        if access_err:
            return access_err
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        result = await get_backup().delete_backup(backup_id)
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool(description="Restore a memory from a base64-encoded archive.")
async def backup_restore_archive(
    archive_base64: Annotated[str, Field(description="Base64-encoded tar.gz archive content")],
    ctx: Optional[Context] = None
) -> dict:
    """
    Restaure une mémoire depuis une archive tar.gz (base64).
    
    L'archive doit contenir manifest.json, graph_data.json, qdrant_vectors.jsonl.
    Si elle contient un dossier documents/, les fichiers sont re-uploadés sur S3.
    
    ⚠️ La mémoire NE DOIT PAS exister (erreur sinon).
    
    Usage typique : backup download --include-documents → fichier.tar.gz → restore-file
    
    Args:
        archive_base64: Contenu de l'archive tar.gz encodé en base64
        
    Returns:
        Compteurs de restauration (entités, relations, vecteurs, documents S3)
    """
    try:
        # Sécurité : vérifier permission write avant restore
        write_err = check_write_permission()
        if write_err:
            return write_err
        
        archive_bytes = base64.b64decode(archive_base64)
        
        # Extraire le memory_id du manifest pour vérifier l'accès
        import tarfile
        from io import BytesIO
        try:
            with tarfile.open(fileobj=BytesIO(archive_bytes), mode='r:gz') as tar:
                manifest_member = tar.getmember('manifest.json')
                manifest_file = tar.extractfile(manifest_member)
                if manifest_file:
                    manifest_data = json.loads(manifest_file.read())
                else:
                    manifest_data = {}
                archive_memory_id = manifest_data.get("memory_id")
                if archive_memory_id:
                    access_err = check_memory_access(archive_memory_id)
                    if access_err:
                        return access_err
        except (KeyError, json.JSONDecodeError):
            pass  # Le backup service gérera les erreurs de format
        
        async def _progress(msg):
            if ctx:
                try:
                    await ctx.info(msg)
                except Exception:
                    pass
        
        result = await get_backup().restore_from_archive(
            archive_bytes=archive_bytes,
            progress_callback=_progress
        )
        return result
    except Exception as e:
        return {"status": "error", "message": str(e)}


# =============================================================================
# Cycle de vie egress (P12-3, Hivemind #268)
# =============================================================================

async def _close_llm_singletons() -> None:
    """Ferme et réinitialise les singletons d'inférence au shutdown.

    Libère les transports proxy possédés (créés uniquement quand PROXY_URL
    est définie) des services extracteur et embedder — les deux registres
    (modules core et caches lazy de ce module) sont réinitialisés. Idempotent
    et no-op quand rien n'a été instancié.
    """
    global _extractor_service, _embedding_service
    from .core.embedder import close_embedding_service_if_initialized
    from .core.extractor import close_extractor_service_if_initialized

    local_extractor = _extractor_service
    local_embedder = _embedding_service
    _extractor_service = None
    _embedding_service = None
    await close_extractor_service_if_initialized()
    await close_embedding_service_if_initialized()
    # Les caches lazy de ce module peuvent référencer les mêmes instances que
    # les singletons de module (déjà fermés ci-dessus) ou des instances
    # injectées : fermer explicitement couvre les deux cas (close idempotent).
    if local_extractor is not None:
        await local_extractor.close()
    if local_embedder is not None:
        await local_embedder.close()


class EgressLifecycleMiddleware:
    """Shim ASGI : ferme les transports proxy possédés au lifespan.shutdown.

    Posé en couche EXTERNE de la pile dans ``main()`` pour que l'arrêt réel
    d'uvicorn (SIGTERM/SIGINT) atteigne toujours le chemin de fermeture,
    quel que soit le comportement des couches internes. Les messages lifespan
    sont transmis inchangés ; les scopes non-lifespan passent tels quels.
    """

    def __init__(self, app):
        self.app = app

    async def __call__(self, scope, receive, send):
        if scope["type"] != "lifespan":
            return await self.app(scope, receive, send)

        async def wrapped_receive():
            message = await receive()
            if message["type"] == "lifespan.shutdown":
                try:
                    await _close_llm_singletons()
                except Exception as e:
                    # Un échec de fermeture ne doit pas casser le shutdown
                    # des couches internes ; message redacted (jamais d'URL
                    # proxy brute).
                    print(
                        "⚠️ [Egress] close on shutdown failed: "
                        f"{redact_proxy_secrets(str(e))}",
                        file=sys.stderr,
                    )
            return message

        await self.app(scope, wrapped_receive, send)


# =============================================================================
# Point d'entrée
# =============================================================================

def main():
    """Point d'entrée principal."""
    parser = argparse.ArgumentParser(description="MCP Memory Server")
    parser.add_argument("--port", type=int, default=settings.mcp_server_port)
    parser.add_argument("--host", type=str, default=settings.mcp_server_host)
    parser.add_argument("--debug", action="store_true", default=settings.mcp_server_debug)
    args = parser.parse_args()
    
    # Récupérer l'app ASGI Streamable HTTP de FastMCP
    # Remplace l'ancien mcp.sse_app() — endpoint unique /mcp au lieu de /sse + /messages
    # Le HostNormalizerMiddleware n'est plus nécessaire (plus de validation Host par Starlette)
    base_app = mcp.streamable_http_app()
    
    # Empiler les middlewares (le dernier wrappé est le premier exécuté)
    # Flux requête : EgressLifecycleMiddleware (lifespan uniquement) →
    # AuthMiddleware → LoggingMiddleware → StaticFilesMiddleware → MCP app
    app = StaticFilesMiddleware(base_app)
    app = LoggingMiddleware(app, debug=args.debug)
    app = AuthMiddleware(app, debug=args.debug)
    # P12-3 : couche externe — ferme les transports proxy possédés des
    # services d'inférence sur lifespan.shutdown (arrêt uvicorn).
    app = EgressLifecycleMiddleware(app)
    
    # Sécurité v2.1.0 : vérifier la clé bootstrap au démarrage
    check_bootstrap_key_safety(settings.admin_bootstrap_key or "")
    
    # Afficher le banner
    print("=" * 70, file=sys.stderr)
    print("🧠 MCP Memory Server - Starting (Streamable HTTP)", file=sys.stderr)
    print(f"📡 Listening on http://{args.host}:{args.port}", file=sys.stderr)
    print(f"🔗 MCP     : http://{args.host}:{args.port}/mcp", file=sys.stderr)
    print(f"🔒 Auth     : Bearer Token (ou ADMIN_BOOTSTRAP_KEY)", file=sys.stderr)
    print(f"🐛 Debug    : {'ENABLED' if args.debug else 'Disabled'}", file=sys.stderr)
    print("=" * 70, file=sys.stderr)
    print("Outils disponibles:", file=sys.stderr)
    print("  - memory_create, memory_delete, memory_list, memory_stats", file=sys.stderr)
    print("  - memory_ingest, memory_search, memory_query, memory_get_context", file=sys.stderr)
    print("  - admin_create_token, admin_list_tokens, admin_revoke_token, admin_update_token", file=sys.stderr)
    print("  - storage_check, storage_cleanup, system_health, system_about, system_whoami", file=sys.stderr)
    print("  - backup_create, backup_list, backup_restore, backup_download, backup_delete", file=sys.stderr)
    print("=" * 70, file=sys.stderr)
    
    # Lancer le serveur
    uvicorn.run(app, host=args.host, port=args.port)


if __name__ == "__main__":
    main()
